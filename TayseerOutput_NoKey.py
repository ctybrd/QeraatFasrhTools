import os
from docx import Document
from docx.shared import Inches, Pt
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from PIL import Image


def read_comments_from_database():
    conn = sqlite3.connect('comments.db')
    c = conn.cursor()
    c.execute(
        "SELECT page_number, id, comment, icon, x_coordinate, y_coordinate FROM comments WHERE (comment IS NOT NULL) AND (comment <> '') AND (icon IN ('/Comment', '/Circle', '/Key')) ORDER BY page_number, y_coordinate desc, x_coordinate desc"
    )
    rows = c.fetchall()
    comments_table = {}
    for row in rows:
        page_number, comment_id, comment, icon, x_coordinate, y_coordinate = row
        comment_info = {
            'comment': comment,
            'icon': icon,
            'x_coordinate': x_coordinate,
            'y_coordinate': y_coordinate,
            'color': get_color_for_icon(icon)  # Replace with your logic to get the color based on the icon
        }
        comments_table.setdefault(page_number, []).append((comment_id, comment_info))
    conn.close()
    return comments_table


def create_word_document(comments_table):
    doc = Document()

    # Set document orientation and page size
    doc.sections[0].orientation = WD_ORIENTATION.PORTRAIT
    doc.sections[0].page_width = Inches(8.27)
    doc.sections[0].page_height = Inches(11.69)

    # Iterate over each page and add the corresponding image, comments, and separator
    for page_number in range(1, 523):
        section = doc.sections[-1] if doc.sections else doc.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        # Get the path of the commented image
        image_path = f'E:/Qeraat/QeraatFasrhTools/Tayseer/Pages_Commented/page_{page_number}_with_comments.png'

        # Check if the image file exists
        if os.path.exists(image_path):
            # Add the image to the document
            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(4))

            # Set paragraph alignment
            if page_number % 2 == 0:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Retrieve comments from the table for the current page
            page_comments = comments_table.get(page_number, [])

            # Add comments to the document
            for i, (comment_id, comment_info) in enumerate(page_comments, start=1):
                comment_text = comment_info['comment']
                icon = comment_info['icon']
                color = comment_info['color']

                if comment_text:  # Add comment only if it's not empty
                    cleaned_comment = comment_text.replace('\n\r', ' ـ ')
                    cleaned_comment = cleaned_comment.replace('\r', ' ـ ')
                    cleaned_comment = cleaned_comment.replace('\n', ' ـ ')

                    # Add the comment to the document
                    paragraph = doc.add_paragraph()
                    run = paragraph.add_run()

                    # Add the icon, comment number, and cleaned comment
                    run.text = get_icon_for_comment(icon) + str(i) + ' ـ ' + cleaned_comment

                    # Set font color
                    run.font.color.rgb = color

                    # Set line spacing
                    paragraph.paragraph_format.line_spacing = Pt(12)

            # Add separator after each page
            doc.add_paragraph()
            doc.add_page_break()

    # Set consistent page margins for all sections
    for section in doc.sections:
        section.left_margin = Inches(1.27)
        section.right_margin = Inches(1.27)
        section.top_margin = Inches(1.27)
        section.bottom_margin = Inches(1.27)

    # Save the document
    doc.save('TayseerOutput.docx')


# Helper function to get the icon representation for a comment
def get_icon_for_comment(icon):
    if icon == '/Comment':
        return 'ش'
    elif icon == '/Circle':
        return 'د'
    elif icon == '/Key':
        return 'ت'
    else:
        return 'م'


# Helper function to get the color based on the icon
def get_color_for_icon(icon):
    if icon == '/Circle':
        return RGBColor(0x00, 0x00, 0xFF)  # Blue color
    elif icon == '/Comment':
        return RGBColor(0x00, 0x00, 0x00)  # Black color
    elif icon == '/Key':
        return RGBColor(0x00, 0x64, 0x00)  # Dark green color
    else:
        return None


# Read comments from the database
comments_table = read_comments_from_database()

# Create the Word document
create_word_document(comments_table)
