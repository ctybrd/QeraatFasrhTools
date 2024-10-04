from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# Folder containing images
image_folder = r"E:\MadinaPages"
image_count = 604

# Create a new Word document
doc = Document()

# Set document size to A3
section = doc.sections[-1]
section.page_width = Inches(16.54)  # A3 width in inches
section.page_height = Inches(11.69)  # A3 height in inches
section.orientation = WD_ORIENT.PORTRAIT  # Optional for horizontal A3

for i in range(1, image_count + 1):
    # Create a new section for each image (optional)
    if i > 1:
        doc.add_page_break()

    # Format the filename with leading zeros
    image_path = os.path.join(image_folder, f"{i:03}.png")
    
    # Add the image and position based on odd/even page
    if i % 2 == 1:  # Odd pages
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:  # Even pages
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(6), height=Inches(8))

# Save the document
doc.save("E:\\MadinaPages\\TadweenMadinaA3.docx")
