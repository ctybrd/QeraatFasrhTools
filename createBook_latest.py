# select page_number,text,group_concat(aya_number) ayas,
# group_concat(ayakey) gkey,
# case when min(aya_number)=max(aya_number) then min(aya_number) else min(aya_number)||' - ' || max(aya_number) end as ayas1
#  from(Select printf('%03d', mosshf_shmrly.sora_number) || printf('%03d', mosshf_shmrly.aya_number)  as ayakey,book_moyassar.aya_index,book_moyassar.text,mosshf_shmrly.page_number,
# mosshf_shmrly.aya_number,mosshf_shmrly.sora_number from book_moyassar
# join mosshf_shmrly on book_moyassar.aya_index=mosshf_shmrly.aya_index
# order by book_moyassar.aya_index)
# group by page_number,text
# order by page_number,gkey

import sqlite3
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from PIL import Image, ImageOps
import re

def connect_to_database(database_path):
    conn = sqlite3.connect(database_path)
    cursor = conn.cursor()
    return conn, cursor

def retrieve_data(cursor):
    query = '''
    SELECT mj.text, ms.page_number, ms.aya_number
    FROM book_qhamza AS mj
    JOIN mosshf_shmrly AS ms ON mj.aya_index = ms.aya_index
    '''
    cursor.execute(query)
    data = cursor.fetchall()
    return data

def process_html_text(html_text):
    soup = BeautifulSoup(html_text, 'html.parser')
    return soup.get_text()

def add_frame_to_image(image_path):
    image = Image.open(image_path)
    framed_image = ImageOps.expand(image, border=10, fill='black')
    framed_image_path = 'framed_image.jpg'
    framed_image.save(framed_image_path)
    return framed_image_path

def create_word_document(data):
    doc = Document()

    # Set consistent page margins for all sections
    for section in doc.sections:
        section.left_margin = Inches(1.27)
        section.right_margin = Inches(1.27)
        section.top_margin = Inches(1.27)
        section.bottom_margin = Inches(1.27)

    current_page = None
    color_switch = False  # Flag to switch between black and blue colors

    for text, page_number, aya_number in data:
        if current_page is None or current_page != page_number:
            if current_page is not None:
                doc.add_page_break()

            current_page = page_number

            image_path = f'e:/pageshamza/{current_page}.png'  # Replace with the actual path to the folder and image file extension

            # Add frame to the image
            framed_image_path = add_frame_to_image(image_path)

            paragraph = doc.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(framed_image_path, width=Inches(4))

            if current_page % 2 == 0:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            comments_paragraph = doc.add_paragraph()
            comments_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        processed_text = process_html_text(text)
        processed_text = processed_text.replace('.', '')
        processed_text = processed_text.replace('\r\n', ' ـ ')
        processed_text = processed_text.replace('\r', ' ـ ')
        processed_text = processed_text.replace('\n', ' ـ ')
        processed_text = 'ـ' + str(aya_number) + 'ـ ' + processed_text+'\n'
        processed_text = re.sub(' {2,}', ' ', processed_text)

        
        run = comments_paragraph.add_run()

        # Set text color to black or blue based on the color_switch flag
        if color_switch:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black color


        run.text = processed_text
        comments_paragraph.paragraph_format.line_spacing = Pt(11)

        color_switch = not color_switch  # Toggle the color_switch flag

    doc.save('output.docx')

def close_database_connection(connection):
    connection.close()

# Main code
database_path = 'e:/qeraat/data_v16.db'
image_folder = 'e:/pageshamza/'

conn, cursor = connect_to_database(database_path)
data = retrieve_data(cursor)
create_word_document(data)
close_database_connection(conn)
