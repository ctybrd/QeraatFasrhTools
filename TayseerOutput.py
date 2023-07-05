import os
from docx import Document
from docx.shared import Inches, Pt
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from PIL import Image, ImageDraw, ImageFont


def read_comments_from_database():
    conn = sqlite3.connect('comments.db')
    c = conn.cursor()
    c.execute(
        "SELECT page_number, id, comment, icon, x_coordinate, y_coordinate FROM comments WHERE (comment IS NOT NULL) AND (comment <> '') AND (icon IN ('/Comment', '/Circle', '/Key')) ORDER BY page_number, y_coordinate desc, x_coordinate desc"
    )
    rows = c.fetchall()
    comments_table = {}
    for row in rows:
        page_number, comment_id, comment, icon, x_coordinate, y_coordinate = row
        comment_info = {
            'comment': comment,
            'icon': icon,
            'x_coordinate': x_coordinate,
            'y_coordinate': y_coordinate,
        }
        comments_table.setdefault(page_number, []).append((comment_id, comment_info))
    conn.close()
    return comments_table


def create_word_document(comments_table):
    doc = Document()

    # Get the list of page images from the folder
    image_folder = './tayseer/Pages/'
    image_files = sorted(os.listdir(image_folder), key=lambda x: int(x.split('.')[0]))

    # Set document orientation and page size
    doc.sections[0].orientation = WD_ORIENTATION.PORTRAIT
    doc.sections[0].page_width = Inches(8.27)
    doc.sections[0].page_height = Inches(11.69)

    # Iterate over each page and add the corresponding image, comments, and separator
    for page_number, image_file in enumerate(image_files, start=1):
        section = doc.sections[-1] if doc.sections else doc.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        # Open the image using PIL
        image_path = os.path.join(image_folder, image_file)
        image = Image.open(image_path)

        # Create a copy of the image for drawing comments
        image_with_comments = image.copy()
        draw = ImageDraw.Draw(image_with_comments)

        # Retrieve comments from the table for the current page
        page_comments = comments_table.get(page_number, [])

        # Add comments to the image copy
        for i, (comment_id, comment_info) in enumerate(page_comments, start=1):
            comment_text = comment_info['comment']
            icon = comment_info['icon']
            x_coordinate = comment_info['x_coordinate']
            y_coordinate = comment_info['y_coordinate']

            if comment_text:  # Add comment only if it's not empty
                cleaned_comment = comment_text.replace('\n\r', ' ـ ')
                cleaned_comment = cleaned_comment.replace('\r', ' ـ ')
                cleaned_comment = cleaned_comment.replace('\n', ' ـ ')

                # Calculate the position to draw the comment code
                font = ImageFont.truetype('arial.ttf', size=18)
                text_width, text_height = draw.textsize(str(i) + get_icon_for_comment(icon), font=font)

                if int(x_coordinate) < 200:
                    text_x = 50
                else:
                    text_x = image_with_comments.width - 50 - text_width

                text_y = (800 - int(y_coordinate) + text_height // 2) * (1669 / 800)

                # Draw the comment code on the image copy
                draw.text((text_x, text_y), str(i) + get_icon_for_comment(icon), font=font, fill=(0, 0, 0))

                # Add the comment text to the document
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                run.text = cleaned_comment
                run.font.color.rgb = get_color_for_icon(icon)

                # Set paragraph alignment
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

                # Set line spacing
                paragraph.paragraph_format.line_spacing = Pt(12)

        # Save the image copy with comments
        image_with_comments.save(f'page_{page_number}_with_comments.png')

        # Add the image with comments to the document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(f'page_{page_number}_with_comments.png', width=Inches(4))

        # Set paragraph alignment
        if page_number % 2 == 0:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add the original image to the document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(4))

        # Add separator after each page
        doc.add_paragraph()
        doc.add_page_break()

    # Set consistent page margins for all sections
    for section in doc.sections:
        section.left_margin = Inches(1.27)
        section.right_margin = Inches(1.27)
        section.top_margin = Inches(1.27)
        section.bottom_margin = Inches(1.27)

    # Save the document
    doc.save('TayseerOutput.docx')


# Helper function to get the icon representation for a comment
def get_icon_for_comment(icon):
    if icon == '/Comment':
        return 'ش'
    elif icon == '/Circle':
        return 'د'
    elif icon == '/Key':
        return 'ت'
    else:
        return 'م'


# Helper function to get the color based on the icon
def get_color_for_icon(icon):
    icon_colors = {
        '/Circle': RGBColor(0x00, 0x00, 0xFF),  # Blue color
        '/Comment': RGBColor(0x00, 0x00, 0x00),  # Black color
        '/Key': RGBColor(0x00, 0x64, 0x00),  # Dark green color
    }
    return icon_colors.get(icon)


# Read comments from the database
comments_table = read_comments_from_database()

# Create the Word document
create_word_document(comments_table)
