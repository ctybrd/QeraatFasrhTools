import os
import shutil
from docx import Document
from docx.shared import Inches, Pt
import sqlite3
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from PIL import Image, ImageDraw, ImageFont
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def create_word_document(comments_table):

    if quran_holder_flag:
    #quran holder database update
        conn = sqlite3.connect('E:\Qeraat\data_v15.db')
        cursor = conn.cursor()
        # Create the table if it doesn't already exist
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS book_tayseer10 (
                aya_index INTEGER NOT NULL,
                text TEXT NOT NULL,
                PRIMARY KEY (aya_index)
            )
        ''')
        # Truncate the table if it exists
        cursor.execute("DELETE FROM book_tayseer10")
    doc = Document()
 
     # Get the list of page images from the folder
    image_folder = './tayseer/Pages' + madina_flag + '/'
    image_files = sorted(os.listdir(image_folder), key=lambda x: int(x.split('.')[0]))

    # Iterate over each page and add the corresponding image, comments, and separator
    for page_number, image_file in enumerate(image_files, start=1):
        if quran_holder_flag:
            #prepare for quran holder app
            aggregated_comment = ''
            # Retrieve the minimum aya_index for the current page from the database
            query = f"SELECT MIN(aya_index) as maya FROM mosshf_shmrly WHERE page_number = {page_number}"
            cursor.execute(query)
            result = cursor.fetchone()
            aya_index = result[0] if result is not None else 0
            if aya_index is None:
                aya_index = 0
            print(page_number,aya_index)

        section = doc.sections[-1] if doc.sections else doc.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)

        # Add image to the page
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(os.path.join(image_folder, image_file), width=Inches(4))

        if page_number % 2 == 0:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Retrieve comments from the table for the current page
        page_comments = comments_table.get(page_number, [])

        # Add comments to the document
        for i, (comment_id, comment_info) in enumerate(page_comments, start=1):
            comment = comment_info['comment']
            icon = comment_info['icon']
            x_coordinate = comment_info['x_coordinate']
            y_coordinate = comment_info['y_coordinate']
            cleaned_comment = comment.replace('\r\n', ' ')
            cleaned_comment = cleaned_comment.replace('\r', ' ')
            cleaned_comment = cleaned_comment.replace('\n', ' ')

            cleaned_comment = cleaned_comment.strip()
            if cleaned_comment:  # Add comment only if it's not empty
                if quran_holder_flag:
                    #to use with quran holder data the whole text for the page
                    # aggregated_comment += get_icon_for_comment(icon)+' - '+ cleaned_comment + '\r\n'
                    aggregated_comment += f'<p> {get_icon_for_comment(icon)} - {cleaned_comment}</p>'
                    # Add the comment text to the document
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                # Set paragraph alignment
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                # Set line spacing
                paragraph.paragraph_format.line_spacing = Pt(12)
                run.text = get_icon_for_comment(icon)+str(i)+' ـ ' + cleaned_comment
                run.font.color.rgb = get_color_for_icon(icon)
                if ((page_number==7) or (page_number==46) or (page_number==263)) and (madina_flag==''):
                    paragraph.paragraph_format.line_spacing = Pt(10)
                else:
                    paragraph.paragraph_format.line_spacing = Pt(12)
                # # Set Sakkal Majalla font for Arabic text
                # arabic_text = run._element
                # arabic_text.rPr.rFonts.set(nsdecls('w:eastAsia'), 'Sakkal Majalla')

        if madina_flag == '':
            if page_number!=522:
                doc.add_page_break()
        if quran_holder_flag:
            query = f"INSERT INTO book_tayseer10 (aya_index, text) VALUES ({aya_index}, '{aggregated_comment}')"
            print (query)
            cursor.execute(query)
            conn.commit()

    # Set consistent page margins for all sections
    for section in doc.sections:
        section.left_margin = Inches(1.27)
        section.right_margin = Inches(1.27)
        section.top_margin = Inches(1.27)
        section.bottom_margin = Inches(1.27)

    if quran_holder_flag:
        cursor.close()
        conn.close()
    # Save the document
    doc.save('TayseerOutput' + madina_flag + '.docx')


def add_comment_keys_to_images(comments_table):
    source_folder = 'E:/Qeraat/QeraatFasrhTools/Tayseer/Pages_Bak' + madina_flag
    image_folder = 'E:/Qeraat/QeraatFasrhTools/Tayseer/Pages' + madina_flag
    for file in os.listdir(image_folder):
        os.remove(os.path.join(image_folder, file))


    # Get the list of page images from the folder
    image_files = sorted(os.listdir(source_folder), key=lambda x: int(x.split('.')[0]))

    # Iterate over each page and add the comment keys to the corresponding image
    for page_number, image_file in enumerate(image_files, start=1):
        # Open the image using PIL
        image_path = os.path.join(source_folder, image_file)
        image = Image.open(image_path)
        draw = ImageDraw.Draw(image)

        # Retrieve comments from the table for the current page
        page_comments = comments_table.get(page_number, [])

        # Add comment keys to the image
        for i, (comment_id, comment_info) in enumerate(page_comments, start=1):
            icon = comment_info['icon']
            x_coordinate = comment_info['x_coordinate']
            y_coordinate = comment_info['y_coordinate']

            # Calculate the position to draw the comment key
            font = ImageFont.truetype('arial.ttf', size=24)
            text = str(i) + get_icon_for_comment(icon)
            text_width, text_height = draw.textsize(text, font=font)

            if int(x_coordinate) < 200:
                text_x = 30
            else:
                text_x = image.width - 30 - text_width

            text_y = (800 - int(y_coordinate) + text_height // 2) * (1669 / 800)

            # Draw the comment key on the image
            draw.text((text_x, text_y), text, font=font, fill=(0, 0, 0))

        # Save the modified image (overwrite the original image)
        modified_image_path = os.path.join(image_folder, f"{page_number}.jpg")
        image.save(modified_image_path)


def read_comments_from_database():
    conn = sqlite3.connect('comments'+madina_flag+'.db')
    c = conn.cursor()
    c.execute(
        "SELECT page_number, id, comment, icon, x_coordinate, y_coordinate FROM comments WHERE (comment IS NOT NULL) AND (comment <> '') AND (icon IN ('/Comment', '/Circle', '/Key')) ORDER BY page_number, y_coordinate desc, x_coordinate desc"
    )
    rows = c.fetchall()
    comments_table = {}
    for row in rows:
        page_number, comment_id, comment, icon, x_coordinate, y_coordinate = row
        comment_info = {
            'comment': comment,
            'icon': icon,
            'x_coordinate': x_coordinate,
            'y_coordinate': y_coordinate,
        }
        comments_table.setdefault(page_number, []).append((comment_id, comment_info))
    conn.close()
    return comments_table


# Helper function to get the icon representation for a comment
def get_icon_for_comment(icon):
    if icon == '/Comment':
        return 'ش'
    elif icon == '/Circle':
        return 'د'
    elif icon == '/Key':
        return 'ت'
    else:
        return 'م'


# Helper function to get the color based on the icon
def get_color_for_icon(icon):
    icon_colors = {
        '/Circle': RGBColor(0x00, 0x00, 0xFF),  # Blue color
        '/Comment': RGBColor(0x00, 0x00, 0x00),  # Black color
        '/Key': RGBColor(0x00, 0x64, 0x00),  # Dark green color
    }
    return icon_colors.get(icon)
# Read comments from the database
quran_holder_flag = False
madina_flag = 'M' # M for mushaf al mdina
comments_table = read_comments_from_database()
# Call the function to add comment keys to the images
add_comment_keys_to_images(comments_table)
# Create the Word document
create_word_document(comments_table)
