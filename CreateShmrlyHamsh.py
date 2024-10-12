import sqlite3
from docx import Document
from docx.shared import Pt
from docx.shared import RGBColor
import os

# Replacement dictionary for reading field
replacement_dict = {
    "بالنقل، مع قصر البدل، وتفخيم الراء.": "بالنقل",
    "بنقل حركة الهمزة.": "بالنقل",
    "بضم ميم الجمع، ووصلها بواو لفظية وصلا.": "صلة ميم الجمع وصلا",
    "بالنقل، مع قصر البدل.": "بالنقل",
    "بالنقل، مع قصر البدل، وتفخيم الراء.": "بالنقل",
    "بالنقل، مع تفخيم الراء.": "بالنقل",
    "بالنقل مع الفتح.": "بالنقل",
    "بالنقل، مع ترك الوقف بهاء السكت.": "بالنقل",
    "بالنقل، مع قصر اللين.": "بالنقل",
    "بالنقل، مع قصر البدل، وترك الوقف بهاء السكت.": "بالنقل",
    "، مع ترك الوقف بهاء السكت":"",
    "، مع تحقيق الهمزتين وصلا ووقفا":"",
    "، مع تحقيق الهمزتين":"",

    # Add more replacements as needed
}

def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
        '(': ')',
        ')': '(',
        '[': ']',
        ']': '[',
        '.' : '٠',                  
    }
    return ''.join(mapping.get(char, char) for char in str(number))

# Connect to the SQLite database
script_path = os.path.abspath(__file__)
drive, _ = os.path.splitdrive(script_path)
drive = drive + '/'
db_path = os.path.join(drive,"Qeraat/QeraatFasrhTools/QeraatSearch/qeraat_data_simple.db")
output_folder = os.path.join(drive,"Qeraat/QeraatFasrhTools/QeraatSearch/output")
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Execute the SQL query
sql_query = """
SELECT aya, sub_subject, reading, page_number2 ,resultnew from quran_data
where 
reading not like 'بصلة ميم الجمع وصلا بخلف.'
and 
reading not like '٠بضم ميم الجمع، ووصلها بواو لفظية بخلف'
and r5_2 is null
and r1_1 is not null
    ORDER BY aya_index, id
"""

cursor.execute(sql_query)

# Create a new Word document
doc = Document()

# Set the page width to 30mm
section = doc.sections[0]
section.page_width = Pt(25 * 28.35)  # converting mm to points (1 mm = 28.35 points)

# Iterate through the SQL results and add them to the document
current_page_number = None
for aya, sub_subject, reading, page_number,resultnew in cursor.fetchall():
    # Add new page for each change in value of page_number2
    if current_page_number is None or current_page_number != page_number:
        doc.add_page_break()
        current_page_number = page_number

    # Add block content to the document
    para = doc.add_paragraph()

    # Sub_subject column value with color based on reading content
    sub_subject_run = para.add_run(transliterate_number(f"{aya} "))
    sub_subject_run.font.size = Pt(10)
    if resultnew:
       sub_subject1 =  
    sub_subject_run = para.add_run(f"{sub_subject}\n")
    # Reading column value with replacements
    edited_text = replacement_dict.get(reading, reading)
    if sub_subject in "'وهو','فهو','لهو','وهو','فهي','لهي','فهي','وهي','ثم هو'":
        sub_subject_run.font.color.rgb = RGBColor(205, 82, 82)  
    elif 'دغ' in reading:
        sub_subject_run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
    else:
        sub_subject_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # Red (default)
    edited_text = edited_text.replace("ـ", "")
    edited_text = transliterate_number('0' + edited_text)  # for Arabic direction
    reading_run = para.add_run(edited_text)
    reading_run.font.size = Pt(11)
    reading_run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black

# Save the document to the specified folder

doc.save(output_folder + "/qaloon_hamsh.docx")

# Close the database connection
conn.close()
