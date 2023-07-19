import sqlite3
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION
import re


def replace_characters(text):
    text = text.replace("ﭼ", ")")
    text = text.replace("ﭽ", "(")
    text = text.replace('.','')
    mapping = {
        '٠': '0',
        '١': '1',
        '٢': '2',
        '٣': '3',
        '٤': '4',
        '٥': '5',
        '٦': '6',
        '٧': '7',
        '٨': '8',
        '٩': '9'
    }
    pattern = re.compile("|".join(map(re.escape, mapping.keys())))
    replaced_text = pattern.sub(lambda match: mapping[match.group(0)], text)
    return replaced_text

def parse_text_and_insert_to_database(text, page_number, cursor):
    matches = re.findall(r"(\d+)(.*?)((?=\d)|$)", text)
    
    for match in matches:
        number = match[0]
        current_text = match[1].strip()
        cursor.execute("INSERT INTO Hamza_Waqf (aya_number, text, page_number) VALUES (?, ?, ?)", (number, current_text, page_number))
    

def extract_text_next_to_keyword(doc_path, keyword):
    conn = sqlite3.connect("Hamza_Waqf_db.db")
    cursor = conn.cursor()
    cursor.execute("CREATE TABLE IF NOT EXISTS Hamza_Waqf (aya_number INTEGER, text TEXT, page_number INTEGER)")
    cursor.execute("DELETE FROM Hamza_Waqf")

    document = Document(doc_path)
    page_number = 3
   
    for table in document.tables:
        incr = False
        for row in table.rows:
            for i in range(len(row.cells)):
                cell = row.cells[i]
                cell_text = cell.text.strip()
                if cell_text == 'السكت' or cell_text == 'الوقف' or cell_text == 'الإمالة':
                    incr = True
                if keyword == cell_text:
                    if i + 1 < len(row.cells):
                        next_cell = row.cells[i + 1]
                        next_cell_text = next_cell.text.strip()
                        replaced_text = replace_characters(next_cell_text)
                        print(page_number, replaced_text)
                        parse_text_and_insert_to_database(replaced_text, page_number, cursor)

        if incr:          
            page_number += 1

    # Query to fill the sora_number column in Hamza_Waqf from mosshf_madina
    update_sora_query = """
    UPDATE Hamza_Waqf
    SET sora_number = (
        SELECT sora_number
        FROM mosshf_madina
        WHERE mosshf_madina.page_number = Hamza_Waqf.page_number
        )
    WHERE EXISTS (
        SELECT sora_number
        FROM mosshf_madina
        WHERE mosshf_madina.page_number = Hamza_Waqf.page_number
        );
    """

    # Execute the update query
    cursor.execute(update_sora_query)
    conn.commit()

    # Query to update spage_number column in Hamza_Waqf from mosshf_shmrly
    update_spage_query = """
    UPDATE Hamza_Waqf
    SET spage_number = (
        SELECT page_number
        FROM mosshf_shmrly
        WHERE mosshf_shmrly.sora_number = Hamza_Waqf.sora_number
        AND mosshf_shmrly.aya_number = Hamza_Waqf.aya_number
        )
    WHERE EXISTS (
        SELECT page_number
        FROM mosshf_shmrly
        WHERE mosshf_shmrly.sora_number = Hamza_Waqf.sora_number
        AND mosshf_shmrly.aya_number = Hamza_Waqf.aya_number
        );
    """
    cursor.execute(update_spage_query)
    update_remaining_spage_query = """
    UPDATE Hamza_Waqf
    SET spage_number = (
        SELECT h2.spage_number
        FROM Hamza_Waqf h2
        WHERE h2.page_number = Hamza_Waqf.page_number
        AND h2.spage_number IS NOT NULL
        ORDER BY h2.page_number DESC
        LIMIT 1
    )
    WHERE spage_number IS NULL;
    """

    # Execute the update query for remaining null values
    cursor.execute(update_remaining_spage_query)
    conn.commit()
    conn.close()


def generate_word_document():
    conn = sqlite3.connect("Hamza_Waqf_db.db")
    cursor = conn.cursor()

    # Retrieve data from the database ordered by spage_number
    cursor.execute("SELECT aya_number, text, spage_number, page_number FROM Hamza_Waqf ORDER BY id")
    rows = cursor.fetchall()

    # Create a new Word document
    document = Document()

    # Iterate through the rows and break pages when spage_number changes
    current_spage_number = None

    for row in rows:
        aya_number, text, spage_number, page_number = row

        # # If spage_number changes, insert a section break
        # if spage_number != current_spage_number:
        #     if current_spage_number is not None:
        #         document.add_page_break()
        #     current_spage_number = spage_number

        # Create a new paragraph for each row of data
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Format the text within brackets as red font for each occurrence
        inside_brackets = False
        run = paragraph.add_run(transliterate_number(aya_number) + " " )
        for char in text:
            if char in ("(",")"):
                run = paragraph.add_run(" ")
            else:
                run = paragraph.add_run(char)
            if char == "(":
                inside_brackets = True
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red font color
            elif char == ")":
                inside_brackets = False
                run.font.color.rgb = RGBColor(0, 0, 0)  # Red font color
            elif inside_brackets:
                run.font.color.rgb = RGBColor(255, 0, 0)  # Red font color

        # Concatenate the transliterated aya number in front of the formatted text
        # formatted_text = text #transliterate_number(aya_number) + " " + text
        # paragraph.add_run(formatted_text)

    # Save the document
    document.save("extracted_data.docx")

    conn.close()

def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩'
    }
    return ''.join(mapping.get(char, char) for char in str(number))


# Provide the path to your Word document and the keyword
doc_path = "E:/Qeraat/QeraatFasrhTools_Data/AliSalehHamza.docx"
keyword = "الوقف"

extract_text_next_to_keyword(doc_path, keyword)

# Call the function to generate the Word document
generate_word_document()
