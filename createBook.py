import sqlite3
from docx import Document
from docx.shared import Inches
from PIL import Image
from bs4 import BeautifulSoup


# Function to convert PNG image to JPEG with white background
def convert_image(image_path, converted_image_path):
    image = Image.open(image_path)
    image = image.convert("RGBA")
    background = Image.new("RGB", image.size, (255, 255, 255))
    background.paste(image, mask=image.split()[3])
    background.save(converted_image_path, "JPEG", quality=100)


# Connect to the SQLite database
conn = sqlite3.connect('e:/qeraat/data_v15.db')
cursor = conn.cursor()

# Execute the query to retrieve data from both tables
query = '''
SELECT mj.text, ms.page_number
FROM book_jlalin AS mj
JOIN mosshf_shmrly AS ms ON mj.aya_index = ms.aya_index
'''
cursor.execute(query)
data = cursor.fetchall()

# Create a new Word document
doc = Document()

# Initialize variables for tracking current page number and concatenated text
current_page = None
concatenated_text = ""

# Iterate over the retrieved data
for text, page_number in data:
    # Check if the page has changed
    if current_page is None or current_page != page_number:
        # Add the concatenated text and image to the document (except for the first iteration)
        if current_page is not None:
            # Add a new paragraph with the concatenated text
            doc.add_paragraph(concatenated_text)

            # Add the image to the right of the text
            image_path = f'E:/Qeraat/pages/{current_page}.png'  # Replace with the actual path to the folder and image file extension
            converted_image_path = f'E:/Qeraat/pages/{current_page}.jpg'  # Replace with the path and filename for the converted image
            convert_image(image_path, converted_image_path)
            doc.add_picture(converted_image_path, width=Inches(3))  # Adjust the width as needed

            # Add a page break for the next page
            doc.add_page_break()

        # Update the current page number and reset the concatenated text
        current_page = page_number
        concatenated_text = ""

    # Process the HTML text and concatenate it
    soup = BeautifulSoup(text, 'html.parser')
    concatenated_text += soup.get_text() + " "  # Modify the separator as needed

# Add the last page's text and image to the document
doc.add_paragraph(concatenated_text)
image_path = f'E:/Qeraat/pages/{current_page}.png'  # Replace with the actual path to the folder and image file extension
converted_image_path = f'E:/Qeraat/pages/{current_page}.jpg'  # Replace with the path and filename for the converted image
convert_image(image_path, converted_image_path)
doc.add_picture(converted_image_path, width=Inches(3))

# Save the Word document
doc.save('output.docx')

# Close the database connection
conn.close()
