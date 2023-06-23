import sqlite3
from docx import Document
from docx.shared import Inches
from PIL import Image
from bs4 import BeautifulSoup


# Function to convert PNG image to JPEG with white background
def convert_image(image_path, converted_image_path):
    image = Image.open(image_path)
    image = image.convert("RGBA")
    background = Image.new("RGB", image.size, (255, 255, 255))
    background.paste(image, mask=image.split()[3])
    background.save(converted_image_path, "JPEG", quality=100)


# Connect to the SQLite database
conn = sqlite3.connect('e:/qeraat/data_v15.db')
cursor = conn.cursor()

# Execute the query to retrieve data from both tables
query = '''
SELECT mj.text, ms.page_number
FROM book_jlalin AS mj
JOIN mosshf_shmrly AS ms ON mj.aya_index = ms.aya_index
'''
cursor.execute(query)
data = cursor.fetchall()

# Create a new Word document
doc = Document()

# Initialize variables for tracking current page number and concatenated text
current_page = None
concatenated_text = ""

# Iterate over the retrieved data
for text, page_number in data:
    # Check if the page has changed
    if current_page is None or current_page != page_number:
        # Add the concatenated text and image to the document (except for the first iteration)
        if current_page is not None:
            # Create a new paragraph for the table
            paragraph = doc.add_paragraph()

            # Create a new table with two columns
            table = doc.add_table(rows=1, cols=2)
            table.autofit = False
            table.columns[0].width = Inches(3)
            table.columns[1].width = Inches(4)

            # Add the image to the left column
            image_path = f'E:/Qeraat/pages/{current_page}.png'  # Replace with the actual path to the folder and image file extension
            converted_image_path = f'E:/Qeraat/pages/{current_page}.jpg'  # Replace with the path and filename for the converted image
            convert_image(image_path, converted_image_path)
            cell_0_0 = table.cell(0, 0)
            cell_0_0.paragraphs[0].alignment = 0  # Left align the image
            run_0_0 = cell_0_0.paragraphs[0].add_run()
            run_0_0.add_picture(converted_image_path, width=Inches(3), height=Inches(4))

            # Add the concatenated text to the right column
            cell_0_1 = table.cell(0, 1)
            cell_0_1.paragraphs[0].alignment = 1  # Center align the text
            run_0_1 = cell_0_1.paragraphs[0].add_run(concatenated_text)

            # Add a page break for the next page
            doc.add_page_break()

        # Update the current page number and reset the concatenated text
        current_page = page_number
        concatenated_text = ""

    # Process the HTML text and concatenate it
    soup = BeautifulSoup(text, 'html.parser')
    concatenated_text += soup.get_text() + " "  # Modify the separator as needed

# Add the last page's text and image to the document
paragraph = doc.add_paragraph()

# Create a new table with two columns
table = doc.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Inches(3)
table.columns[1].width = Inches(4)

# Add the image to the left column
image_path = f'E:/Qeraat/pages/{current_page}.png'  # Replace with the actual path to the folder and image file extension
converted_image_path = f'E:/Qeraat/pages/{current_page}.jpg'  # Replace with the path and filename for the converted image
convert_image(image_path, converted_image_path)
cell_0_0 = table.cell(0, 0)
cell_0_0.paragraphs[0].alignment = 0  # Left align the image
run_0_0 = cell_0_0.paragraphs[0].add_run()
run_0_0.add_picture(converted_image_path, width=Inches(3), height=Inches(4))

# Add the concatenated text to the right column
cell_0_1 = table.cell(0, 1)
cell_0_1.paragraphs[0].alignment = 1  # Center align the text
run_0_1 = cell_0_1.paragraphs[0].add_run(concatenated_text)

# Save the Word document
doc.save('output.docx')

# Close the database connection
conn.close()
