from docx import Document
from docx.shared import Pt, Inches
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
import re
import sqlite3

def create_word_document(comments_table):
    # Create a Word document with custom page dimensions suitable for a phone screen
    doc = Document()
    page_width = Inches(3.07)
    page_height = Inches(6.33)

    for page_number in range(1, len(comments_table) + 1):
        section = doc.sections[-1] if doc.sections else doc.add_section()
        section.orientation = WD_ORIENTATION.PORTRAIT
        section.page_width = page_width
        section.page_height = page_height

        # Retrieve comments for the current page
        page_comments = comments_table.get(page_number, [])

        comment_paragraph = doc.add_paragraph()
        comment_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        comment_paragraph.paragraph_format.line_spacing = Pt(10)  # Set line spacing to 8 points
        comment_paragraph.paragraph_format.space_before = Pt(0)  # Set before paragraph spacing to zero
        comment_paragraph.paragraph_format.space_after = Pt(0)  # Set after paragraph spacing to zero

        for i, (comment_id, comment_info) in enumerate(page_comments, start=1):
            comment = comment_info['comment']
            icon = comment_info['icon']

            cleaned_comment = comment.replace('\r\n', ' ')
            cleaned_comment = cleaned_comment.replace('\r', ' ')
            cleaned_comment = cleaned_comment.replace('\n', ' ')
            cleaned_comment = re.sub(' {2,}', ' ', cleaned_comment)
            cleaned_comment = cleaned_comment.strip()

            if cleaned_comment:
                run = comment_paragraph.add_run()
                run.text = get_icon_for_comment(icon) + str(i) + ' ـ ' + cleaned_comment + ' '  # Add a space after cleaned_comment
                run.font.color.rgb = get_color_for_icon(icon)
                run.font.size = Pt(8)  # Set the font size to 8 points

                # Add a red asterisk as a separate run with red color
                asterisk_run = comment_paragraph.add_run(' ـ ')
                asterisk_run.font.color.rgb = RGBColor(255, 0, 0)  # Red color for the asterisk

        if (page_number != 522) or (madina_flag == 'M'):
            doc.add_page_break()

    # Set consistent page margins for all sections
    for section in doc.sections:
        section.left_margin = Inches(0.2)  # Adjust margins as needed
        section.right_margin = Inches(0.2)
        section.top_margin = Inches(0.2)
        section.bottom_margin = Inches(0.2)

    # Save the document
    doc.save('TayseerOutput' + madina_flag + '.docx')




def read_comments_from_database():
    conn = sqlite3.connect('comments' + madina_flag + '.db')
    c = conn.cursor()
    c.execute(
        "SELECT page_number, id, comment, icon FROM comments WHERE (comment IS NOT NULL) AND (comment <> '') AND (icon IN ('/Comment', '/Circle', '/Key')) ORDER BY page_number"
    )
    rows = c.fetchall()
    comments_table = {}
    for row in rows:
        page_number, comment_id, comment, icon = row
        comment_info = {
            'comment': comment,
            'icon': icon
        }
        comments_table.setdefault(page_number, []).append((comment_id, comment_info))
    conn.close()
    return comments_table

def get_icon_for_comment(icon):
    if icon == '/Comment':
        return 'ش'
    elif icon == '/Circle':
        return 'د'
    elif icon == '/Key':
        return 'ت'
    else:
        return 'م'

def get_color_for_icon(icon):
    icon_colors = {
        '/Circle': RGBColor(0x00, 0x00, 0xFF),  # Blue color
        '/Comment': RGBColor(0x00, 0x00, 0x00),  # Black color
        '/Key': RGBColor(0x00, 0x64, 0x00),  # Dark green color
    }
    return icon_colors.get(icon)

madina_flag = ''  # M for mushaf al mdina, empty for shamarly

# Read comments from the database
comments_table = read_comments_from_database()

# Create the Word document with custom page dimensions
create_word_document(comments_table)
