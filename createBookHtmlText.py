import sqlite3
from docx import Document
from docx.shared import Inches
from PIL import Image
from bs4 import BeautifulSoup

# Function to convert PNG image to JPEG with white background
def convert_image(image_path, converted_image_path):
    image = Image.open(image_path)
    image = image.convert("RGBA")
    background = Image.new("RGB", image.size, (255, 255, 255))
    background.paste(image, mask=image.split()[3])
    background.save(converted_image_path, "JPEG", quality=100)

# Connect to the SQLite database
conn = sqlite3.connect('e:/qeraat/data_v15.db')
cursor = conn.cursor()

# Execute the query to retrieve data from both tables
query = '''
SELECT mj.text, ms.page_number
FROM book_jlalin AS mj
JOIN mosshf_shmrly AS ms ON mj.aya_index = ms.aya_index
'''
cursor.execute(query)
data = cursor.fetchall()

# Create a new Word document
doc = Document()

# Iterate over the retrieved data
for text, page_number in data:
    # Add the image to the document
    image_path = f'E:/Qeraat/pages/{page_number}.png'
    converted_image_path = f'E:/Qeraat/pages/{page_number}.jpg'
    convert_image(image_path, converted_image_path)
    doc.add_picture(converted_image_path, width=Inches(3))

    # Process the HTML text and add it to the document
    soup = BeautifulSoup(text, 'html.parser')
    plain_text = soup.get_text()
    doc.add_paragraph(plain_text)

    # Add a page break for the next page
    doc.add_page_break()

# Save the Word document
doc.save('output.docx')

# Close the database connection
conn.close()
