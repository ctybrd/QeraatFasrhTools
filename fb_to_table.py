from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os
import sqlite3
import re

# Define the add_border function
def add_border(paragraph, color):
    """Adds a rounded border and background color to a paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()

    # Create a shading element for the specified background color
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')  # Type of shading (clear = background fill only)
    shd.set(qn('w:fill'), color)  # Background color
    shd.set(qn('w:color'), 'auto')  # Use automatic text color
    pPr.append(shd)

# Set RTL orientation and narrow margins
from docx.shared import Cm

def configure_doc_settings(doc):
    """Sets document orientation to RTL and margins to 1.25 cm."""
    # Enable RTL for the document
    sectPr = doc.sections[0]._sectPr
    rtl = OxmlElement('w:bidi')
    rtl.set(qn('w:val'), '1')  # Enable RTL layout
    sectPr.append(rtl)

    # Set margins using the `Cm` utility
    section = doc.sections[0]
    section.top_margin = Cm(1.25)
    section.bottom_margin = Cm(1.25)
    section.left_margin = Cm(1.25)
    section.right_margin = Cm(1.25)


# Colors for alternating background
colors = ['E6F7FF', 'FFE6E6']  # Light blue and light red

# Connect to SQLite and query hashtags
conn = sqlite3.connect('facebook_posts.db')
cursor = conn.cursor()

cursor.execute("SELECT DISTINCT hashtags FROM posts")
all_hashtags = set()
for row in cursor.fetchall():
    hashtags = row[0].split(', ')
    all_hashtags.update(hashtags)

output_folder = 'output_documents'
os.makedirs(output_folder, exist_ok=True)

print("Generating Word documents...")

for hashtag in all_hashtags:
    # Create a new Word document
    doc = Document()
    configure_doc_settings(doc)
    doc.add_heading(f'Posts for #{hashtag}', level=1)

    # Query posts with this hashtag
    cursor.execute("SELECT real_datetime, post_text, hashtags FROM posts WHERE hashtags LIKE ?", (f'%{hashtag}%',))
    posts = cursor.fetchall()

    if not posts:
        continue

    # Alternate colors for each post
    for index, (real_datetime, post_text, hashtags) in enumerate(posts):
        if len(hashtags.split(', ')) > 10:
            continue  # Skip if more than 10 hashtags

        # Create a paragraph for each post with enhanced appearance
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"{real_datetime}\n{post_text}")
        run.font.size = Pt(28)
        run.font.name = "Sakkal Majalla"
        run.font.color.rgb = RGBColor(0, 0, 0)

        # Apply border and background color, alternating colors
        add_border(p, colors[index % 2])

        doc.add_paragraph()  # Add a blank line between posts

    # Save the document for this hashtag in the output folder
    sanitized_filename = re.sub(r'[<>:"/\\|?*]', '_', f'hashtag_{hashtag}.docx')
    filename = os.path.join(output_folder, sanitized_filename)
    doc.save(filename)
    print(f"Document created: {filename}")

# Close the database connection
conn.close()

print("All Word documents have been successfully generated in the 'output_documents' folder.")
