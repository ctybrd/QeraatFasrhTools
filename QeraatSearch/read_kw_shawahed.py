import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement, ns

def cmyk_to_rgb(c, m, y, k):
    """Convert CMYK values (0 to 1 range) to RGB values (0 to 255 range)."""
    r = 255 * (1 - c) * (1 - k)
    g = 255 * (1 - m) * (1 - k)
    b = 255 * (1 - y) * (1 - k)
    return int(r), int(g), int(b)

def extract_unicode(char_data):
    """Extract unicode text from character data."""
    try:
        return char_data['unicode']
    except (KeyError, TypeError):
        return None

def process_char_data(char_data, paragraph):
    """Process individual character data and add it to the Word document."""
    text = extract_unicode(char_data)

    if not text:
        return

    if text == "*" or text[0] == '-':  # Break before '*'
        paragraph = doc.add_paragraph()  # Start a new paragraph

    text = text.replace(')', '<b>').replace('(', '</b>')
    text = text.replace('<b>', '(').replace('</b>', ')')
    text = text.replace('-', 'ـ')

    # Handle text direction if 'upright' is specified
    if char_data.get('upright'):
        run = paragraph.add_run(text)
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.text = '1'  # '1' for right-to-left, '0' for left-to-right
        rPr.append(rtl)
    else:
        run = paragraph.add_run(text)
    
    # Set font name
    fontname = char_data.get('fontname')
    if fontname:
        run.font.name = fontname
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(ns.qn('w:ascii'), fontname)
        rFonts.set(ns.qn('w:hAnsi'), fontname)
        rFonts.set(ns.qn('w:cs'), fontname)

    # Set font color
    color = char_data.get('color')
    if isinstance(color, dict) and 'ncolor' in color:
        cmyk_values = color['ncolor']
        if len(cmyk_values) == 4:
            r, g, b = cmyk_to_rgb(*cmyk_values)
            run.font.color.rgb = RGBColor(r, g, b)

    # Set font size (convert to Pt)
    size = char_data.get('size')
    if size:
        run.font.size = Pt(size)

    # Add a line break after a period character
    if text.endswith('.'):
        paragraph.add_run().add_break()

def create_document(data, pagenumber, doc):
    """Creates the content in the Word document based on the JSON data."""
    doc.add_paragraph(f"صفحة: {pagenumber}").bold = True

    for item in data:
        paragraph = doc.add_paragraph()  # Initialize paragraph at the start of each item

        # Process 'shahed_chars'
        shahed_chars = item.get('shahed_chars', [])
        if isinstance(shahed_chars, list):
            for char_list in shahed_chars:
                if isinstance(char_list, list):
                    for char_data in char_list:
                        if isinstance(char_data, dict):
                            process_char_data(char_data, paragraph)

        # Process 'dalal_chars'
        dalal_chars = item.get('dalal_chars', [])
        if isinstance(dalal_chars, list):
            for char_list in dalal_chars:
                if isinstance(char_list, list):
                    for char_data in char_list:
                        if isinstance(char_data, dict):
                            process_char_data(char_data, paragraph)

def main():
    """Reads JSON data from multiple files in a folder and generates a Word document with the formatted text."""
    script_path = os.path.abspath(__file__)
    drive, _ = os.path.splitdrive(script_path) 
    drive = drive + '/'
    folder_path = os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json')
    
    # Create a new Word document
    doc = Document()

    try:
        # Extract numeric parts of filenames and sort them
        files = sorted(
            [f for f in os.listdir(folder_path) if f.startswith("Shawahed_") and f.endswith(".json")],
            key=lambda x: int(x.split("_")[1].split(".")[0])
        )

        for filename in files:
            pagenumber = int(filename.split("_")[1].split(".")[0])  # Extract pagenumber from filename
            with open(os.path.join(folder_path, filename), 'r', encoding='utf-8') as f:
                data = json.load(f)
                create_document(data, pagenumber, doc)

                # Insert a page break after processing each file
                doc.add_page_break()

        # Save the Word document
        doc.save('Shawahed_Content.docx')

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == '__main__':
    main()
