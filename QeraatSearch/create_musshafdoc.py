import sqlite3
import re
from docx import Document
from docx.shared import Pt

# File paths
db_path = r"D:\Qeraat\QeraatFasrhTools\qeraat_data_simple.db"
docx_file_path = r"D:\Qeraat\QeraatFasrhTools\Quran_Txt.docx"
output_file_path = r"D:\Qeraat\QeraatFasrhTools\quran_with_breaks.docx"

# Arabic-Indic numerals and special words to ignore in counting
excluded_words = {'۞', '۩'}
arabic_numerals_pattern = r'[١٢٣٤٥٦٧٨٩٠]+'  # Regular expression to match Arabic numerals (1 or more)

# Connect to the database
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Fetch words data from the database
cursor.execute("""
    SELECT wordindex, word, page_number2, lineno2
    FROM words1
    WHERE wordindex NOT IN (1, 2, 3, 4, 30, 31, 32, 33)
    ORDER BY wordindex
""")
words_data = cursor.fetchall()
conn.close()

# Read the input Word document
doc = Document(docx_file_path)

# Initialize variables
current_page = None
current_line = None
word_count = 0
line_words = []  # List to hold words for the current line

# Function to check if the word is an excluded word or contains Arabic numerals
def is_excluded(word):
    if word in excluded_words:
        return True
    if re.match(arabic_numerals_pattern, word):
        return True
    return False

# Extract all text from the document into a list of words
text_words = []
for paragraph in doc.paragraphs:
    text_words.extend(paragraph.text.split())  # Split text into words and add to the list

# Create a new document for the output
output_doc = Document()

# Process words from the database and map them to the text words
surah_started = False  # Flag to track if we are in the middle of a surah

for word_data in words_data:
    wordindex, db_word, page_number2, lineno2 = word_data

    while word_count < len(text_words):
        current_word = text_words[word_count]

        # Handle Surah headers (only once for the first word of surah)
        if current_word.startswith('سُورَةُ') and not surah_started:
            output_doc.add_paragraph()  # Add an empty paragraph before surah header
            surah_header = ' '.join(text_words[word_count:word_count + 2])  # Surah header
            output_doc.add_paragraph(surah_header)  # Add Surah header
            bismillah = ' '.join(text_words[word_count + 2:word_count + 6])  # Bismillah
            output_doc.add_paragraph(bismillah)  # Add Bismillah
            word_count += 6  # Skip surah header and Bismillah
            line_words = []  # Reset line words
            surah_started = True  # Mark surah as started
            current_line = None  # Reset line tracking for the surah
            current_page = None  # Reset page tracking for the surah
            continue  # Move to the next word after the surah header

        # Skip excluded words
        if is_excluded(current_word):
            line_words.append(current_word)
            word_count += 1
            continue

        # Add word to the current line
        line_words.append(current_word)
        word_count += 1

        # Check if the line number changes
        if lineno2 != current_line:
            if line_words:
                # Add the previous line to the document
                run = output_doc.add_paragraph().add_run(' '.join(line_words))
                run.font.name = 'KFGQPC HAFS Uthmanic Script'
                run.font.size = Pt(22)  # Set the font size
                line_words = []  # Reset line words
            current_line = lineno2

        # Check if the page number changes
        if page_number2 != current_page:
            if current_page is not None:  # Avoid adding a page break at the start
                output_doc.add_paragraph().add_run("\f")  # Add page break in Word
            current_page = page_number2

# Add any remaining words in the last line
if line_words:
    run = output_doc.add_paragraph().add_run(' '.join(line_words))
    run.font.name = 'KFGQPC HAFS Uthmanic Script'
    run.font.size = Pt(22)  # Set the font size

# Save the modified document
output_doc.save(output_file_path)
print(f"Output written to: {output_file_path}")
