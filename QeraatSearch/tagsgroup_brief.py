import sqlite3
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, RGBColor
from docx.shared import Pt

# Connect to the SQLite database
conn = sqlite3.connect('E:\\Qeraat\\QeraatFasrhTools\\QeraatSearch\\qeraat_data_simple.db')
cursor = conn.cursor()

# Updated SQL query with GROUP BY
query = """WITH RECURSIVE tags_split(aya_index, id, sub_subject, reading, qareesrest,rasaya, tag, remaining_tags) AS (
    SELECT aya_index, id, sub_subject1 sub_subject, reading, qareesrest,rasaya,
           TRIM(',' || CASE 
               WHEN INSTR(TRIM(tags), ',') > 0 THEN SUBSTR(TRIM(tags), 1, INSTR(TRIM(tags), ',') - 1)
               ELSE TRIM(tags)
           END, ','),
           TRIM(',' || CASE 
               WHEN INSTR(TRIM(tags), ',') > 0 THEN SUBSTR(TRIM(tags), INSTR(TRIM(tags), ',') + 1)
               ELSE ''
           END, ',')
    FROM quran_data 
    WHERE tags IS NOT NULL AND tags NOT LIKE '%nochange%' AND page_shmrly IS NOT NULL

    UNION ALL

    SELECT aya_index, id, sub_subject, reading, qareesrest,rasaya,
           TRIM(',' || CASE 
               WHEN INSTR(TRIM(remaining_tags), ',') > 0 THEN SUBSTR(TRIM(remaining_tags), 1, INSTR(TRIM(remaining_tags), ',') - 1)
               ELSE TRIM(remaining_tags)
           END, ','),
           TRIM(',' || CASE 
               WHEN INSTR(TRIM(remaining_tags), ',') > 0 THEN SUBSTR(TRIM(remaining_tags), INSTR(TRIM(remaining_tags), ',') + 1)
               ELSE ''
           END, ',')
    FROM tags_split
    WHERE remaining_tags != ''
)

SELECT page_shmrly, description, group_concat(distinct sub_subject) AS sub_subject, 
       '' qareesrest, srt,case when rasaya is NOT null then 'رأس آية'
	   ELSE '' END rasaya
FROM (
    SELECT ts.aya_index, 
           ts.id, 
           ts.sub_subject, 
           ts.reading, 
           ts.tag, 
           tm.description,
           tm.srt,
           qd.page_shmrly,
           qd.qareesrest,
		   case when ts.tag in ('imala','taklel') then qd.rasaya else 1 end rasaya
    FROM tags_split ts
    LEFT JOIN tagsmaster tm ON ts.tag = tm.tag
    JOIN quran_data qd ON ts.aya_index = qd.aya_index AND ts.id = qd.id
    WHERE ts.tag != '' AND ts.tag NOT IN ('farsh', 'nakl', 'meemsela', 'sakt', 'saktharf','hoomoo','imalah','haasakt')
    and qd.r1_2 is not null
    ORDER BY qd.page_shmrly, tm.srt, ts.tag,qd.rasaya desc,ts.aya_index, ts.id
)
GROUP BY page_shmrly, srt, tag,rasaya, description
ORDER by  page_shmrly, srt, tag,rasaya desc, description;
"""

# Fetch data
data = cursor.execute(query).fetchall()
conn.close()

# Variables to manage splitting documents
doc = Document()
file_index = 1
page_count = 0
last_page_shmrly = None

# Function to set margins and add page number header
def setup_document():
    section = doc.sections[0]
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.7)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)
    doc.add_heading(f'صفحة: {last_page_shmrly}', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Function to save the document and reset for a new one
def save_and_reset_document():
    global doc, file_index, page_count
    doc.save(f'E:/Qeraat/مشروع العشر/shmrly_osoul_brief_part_{file_index}.docx')
    doc = Document()
    file_index += 1
    page_count = 0
    setup_document()

# Set up the first document
setup_document()

# Generate the document
for row in data:
    page_shmrly = row[0]  # page_shmrly
    description = row[1]   # description
    # reading = row[2]       # reading
    sub_subjects = row[2]  # sub_subject
    qareesrests = row[3]   # qareesrest
    tagsrt = row[4]

    # Check if a new document should be started
    if page_count >= 10:
        save_and_reset_document()

    # Check if we need a page break
    if last_page_shmrly is None or page_shmrly != last_page_shmrly:
        if last_page_shmrly is not None:  # Add a page break only if it's not the first page
            doc.add_page_break()
        last_page_shmrly = page_shmrly
        page_count += 1
        
        # Add a header with the page number
        doc.add_heading(f'صفحة: {page_shmrly}', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add tag description
    doc.add_heading(description, level=2)

    # Create a table for the records
    table = doc.add_table(rows=1, cols=4)
    row_cells = table.add_row().cells
    cell_sub_subject = row_cells[3]
    cell_sub_subject.text = str(sub_subjects)
    for paragraph in cell_sub_subject.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 0, 50)
            run.font.size = Pt(14) 
            row_cells[1].text = '' #str(reading)
    row_cells[0].text = str(qareesrests)

    # Reduce the space between rows
    for cell in row_cells:
        paragraph = cell.paragraphs[0]
        paragraph.paragraph_format.space_after = Cm(0)
        paragraph.paragraph_format.space_before = Cm(0)

    doc.add_paragraph()

# Save the last document if it contains data
if page_count > 0:
    save_and_reset_document()
