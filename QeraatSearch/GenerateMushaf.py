import sqlite3
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm

def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
    }
    return ''.join(mapping.get(char, char) for char in str(number))

# Connect to the SQLite database
conn = sqlite3.connect('E:/Qeraat/QeraatFasrhTools/QeraatSearch/qeraat_data_simple.db')
cursor = conn.cursor()

# Execute the SQL query
query = """
SELECT 
    page_number2,
    GROUP_CONCAT(csub_subject, ', ') as csub_subject,
    creading,
    readingresult,
    GROUP_CONCAT(ayas, ', ') as ayas
FROM (
    SELECT 
        page_number2,
        aya || ' ـ ' || sub_subject || 
        CASE 
            WHEN count_words = 1 THEN '' 
            WHEN count_words = 2 THEN ' )معا(' 
            ELSE ' )جميعا(' 
        END as csub_subject,
        CASE 
            WHEN q6 IS NOT NULL THEN '' 
            ELSE 
                CASE 
                    WHEN r6_1 IS NOT NULL THEN 'خلف ' 
                    ELSE 'خلاد ' 
                END 
        END || reading as creading,
        ifnull(readingresult, '') as readingresult,
        printf('%03d', sora) || printf('%03d', aya) || printf('%03d', id) as ayas 
    FROM quran_data 
    WHERE 
        qareesrest LIKE '%حمزة%' and IFNULL(r5_2, 0) = 0 
        and sub_sno = 1
    ORDER BY page_number2, aya, id
) AS subquery
GROUP BY page_number2, creading, readingresult
ORDER BY page_number2, ayas;

"""

cursor.execute(query)
rows = cursor.fetchall()

# Close the database connection
conn.close()

# Create a new Word document
doc = Document()

# Set document margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.7)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)

# Set font size for the entire document
doc.styles['Normal'].font.size = Pt(13)

# Iterate through all available image files
for page_number2 in range(1, 523):  # Assuming you have 522 pages
    image_path = f'e:/pageshamza/{page_number2}.png'
    
    # Check if the image file exists
    if os.path.exists(image_path):
        # Add the image to the Word document
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        run.add_picture(image_path, width=Cm(11.97))  # A4 width is approximately 13.97 cm
        if page_number2 % 2 == 0:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        else:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Check if there is a corresponding row in the query
        matching_rows = [row for row in rows if row[0] == page_number2]

        # If there are matching rows, add text with different colors
        if matching_rows:
            # we need this if all at one line
            # paragraph = doc.add_paragraph()
            # paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            is_black = True  
            for row in matching_rows:
            # remove this if all at one line
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT               
                # Add Aya and Sub_subject with color
                run = paragraph.add_run()
                run.font.color.rgb = RGBColor(139, 0, 0) #RGBColor(255, 0, 0)
                formatted_aya = transliterate_number(row[1] + ' ')
                run.add_text(formatted_aya)
                run.bold = True  # Make the red part bold
                # run.font.size = Pt(13)  # Set the font size
                # Add Reading with color
                run = paragraph.add_run(' ')
                # remark this if we want alternation colors
                run.font.color.rgb = RGBColor(0, 0, 0)

                # if is_black:
                #     run.font.color.rgb = RGBColor(0, 0, 0)
                # else:
                #     run.font.color.rgb = RGBColor(0, 0, 255)
                # is_black = not is_black 
                readingclean  = row[2].replace(')',' ').replace('(',' ').replace('.','')
                run.add_text(readingclean   )

        # Add a page break after each page_number2 change
        doc.add_page_break()

# Save the Word document
doc.save('e:/pageshamza/Hamzah-Shamrly-Shalaby_New.docx')
