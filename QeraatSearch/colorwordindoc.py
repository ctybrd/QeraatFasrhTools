from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def change_font_color_preserve_formatting(doc_path, search_text, new_color):
    # Load the Word document
    doc = Document(doc_path)
    
    # Define the new color (RGB format)
    color = RGBColor(*new_color)
    
    # Function to change font color while preserving formatting
    def modify_run(run, text, color):
        # Split the run text to isolate the search text
        parts = run.text.split(text)
        if len(parts) > 1:
            # Clear the run text
            run.text = ""
            # Add text before the search text
            run.add_text(parts[0])
            # Add the search text with the new color
            new_run = paragraph.add_run(text)
            # Copy font properties from the original run
            new_run.font.name = run.font.name
            new_run.font.size = run.font.size
            new_run.font.bold = run.font.bold
            new_run.font.italic = run.font.italic
            new_run.font.underline = run.font.underline
            new_run.font.color.rgb = color
            # Add text after the search text
            run.add_text(parts[1])
        else:
            # If the run contains only the search text, change its color
            run.font.color.rgb = color
    
    # Iterate through all paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the search text is in the paragraph
        if search_text in paragraph.text:
            # Set paragraph direction to RTL
            paragraph._element.pPr.add_child(OxmlElement('w:bidi'))
            paragraph._element.pPr.bidi.val = '1'
            # Iterate through all runs in the paragraph
            for run in paragraph.runs:
                if search_text in run.text:
                    modify_run(run, search_text, color)
    
    # Iterate through shapes (e.g., text boxes) in the document
    for shape in doc.inline_shapes:
        if hasattr(shape, 'text_frame'):
            for paragraph in shape.text_frame.paragraphs:
                if search_text in paragraph.text:
                    # Set paragraph direction to RTL
                        # Enable RTL for the document
                    sectPr = doc.sections[0]._sectPr
                    rtl = OxmlElement('w:bidi')
    rtl.set(qn('w:val'), '1')  # Enable RTL layout
    sectPr.append(rtl)
                    paragraph._element.pPr.add_child(OxmlElement('w:bidi'))
                    paragraph._element.pPr.bidi.val = '1'
                    # Iterate through all runs in the paragraph
                    for run in paragraph.runs:
                        if search_text in run.text:
                            modify_run(run, search_text, color)
    
    # Save the modified document
    doc.save("modified_document.docx")
    print(f"Document saved as 'modified_document.docx' with '{search_text}' in new color.")

# Example usage
doc_path = "E:/Qeraat/QeraatFasrhTools/ShmrlyWord/test.docx"  # Path to your Word document
search_text = "مَٰلِكِ"          # Text to search for
new_color = (255, 0, 0)          # New font color in RGB (red in this case)

change_font_color_preserve_formatting(doc_path, search_text, new_color)