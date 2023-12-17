import sqlite3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml import OxmlElement

def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
        ')': '(',
        '(': ')'
    }
    return ''.join(mapping.get(char, char) for char in str(number))

# Define colors as RGBColor objects
RED = RGBColor(255, 0, 0)
GREEN = RGBColor(0, 128, 0)
BLUE = RGBColor(0, 0, 255)

# Open the SQLite database
db_path = "E:\Qeraat\QeraatFasrhTools\QeraatSearch\qeraat_data.db"
connection = sqlite3.connect(db_path)
cursor = connection.cursor()

# Execute the SQLite query
query = """
SELECT  reading,group_concat(sora_name ||' ـ ' ||aya || ' ـ ' ||sub_subject || ' ')  as words
FROM all_qeraat 
WHERE  
    (

(q4 is not null or r4_1 is not null or r4_2 is not null) 
AND (r5_1 is not null)
and (r5_2 is null)
   )
   GROUP BY reading
    order by reading;

"""
cursor.execute(query)

# Create a new Word document for the first sura
doc = None
current_sura = None
current_sura_no = None
current_aya = None
table = None  # Initialize the table to None
color_index = 0  # Index to cycle through colors

doc = Document()
for row in cursor.fetchall():
    reading, words = row    
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    for row in table.rows:
        row = table.add_row().cells
        content_list = [transliterate_number(reading),transliterate_number(words)]
        for i, content in enumerate(content_list):
            p = row[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(content)
            run.bold = True
            run.font.color.rgb = (RED, GREEN, BLUE)[i]

# Save the last sura document
if doc:
    doc.save(f"./output/IbnAmerShoba_brief.docx")
# Close the database connection
connection.close()
