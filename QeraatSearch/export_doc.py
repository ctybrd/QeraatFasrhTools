import sqlite3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩'
    }
    return ''.join(mapping.get(char, char) for char in str(number))


# Open the SQLite database
db_path = "E:\Qeraat\QeraatFasrhTools\QeraatSearch\qeraat_data.db"
connection = sqlite3.connect(db_path)
cursor = connection.cursor()

# Execute the SQLite query
query = """
SELECT sora_name, aya, text_full, sub_subject, qareesrest, reading, sora
FROM all_qeraat
WHERE (Q6=1 OR R6_1=1 OR R6_2=1) AND R5_2 IS NULL
ORDER BY sora, aya, id
"""
cursor.execute(query)

# Create a new Word document for the first sura
doc = None
current_sura = None
current_sura_no = None
current_aya = None
table = None  # Initialize the table to None

for row in cursor.fetchall():
    sora_name, aya, text_full, sub_subject, qareesrest, reading, sora = row

    if sora_name != current_sura:
        # Create a new document for the new sura
        if doc:
            doc.save(f"./output/{current_sura_no}{current_sura}.docx")
        doc = Document()
        current_sura = sora_name
        current_sura_no = sora

    if aya != current_aya:
        # Create a new group header when the aya changes
        if current_aya:
            doc.add_paragraph()  # Add an empty line
        current_aya = aya
        aya_text = transliterate_number(aya)+f" ـ {text_full}"
        para = doc.add_paragraph(aya_text)
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = para.runs[0]
        run.bold = True

        table = doc.add_table(rows=1, cols=3)
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    for run in paragraph.runs:
                        run.bold = True

        # Set borders for the table cells
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        for row in table.rows:
            for cell in row.cells:
                for element in cell._element.iter():
                    if element.tag.endswith('tcBorders'):
                        for border in element:
                            if border.tag.endswith('start'):
                                border.attrib['w:val'] = 'single'
                            if border.tag.endswith('top'):
                                border.attrib['w:val'] = 'single'
                            if border.tag.endswith('end'):
                                border.attrib['w:val'] = 'single'
                            if border.tag.endswith('bottom'):
                                border.attrib['w:val'] = 'single'
                            border.attrib['w:sz'] = '8'
                            border.attrib['w:space'] = '0'
                            border.attrib['w:color'] = 'auto'

    # Add the data to the current table
    row = table.add_row().cells
    for i, content in enumerate([reading, qareesrest, sub_subject]):
        p = row[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(content)
        run.bold = True

# Save the last sura document
if doc:
    doc.save(f"./output/{current_sura_no}{current_sura}.docx")

# Close the database connection
connection.close()
