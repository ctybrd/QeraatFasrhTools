import sqlite3
import requests
from bs4 import BeautifulSoup
from docx import Document

site = 'ibn-atiyah'

# Connect to the SQLite database
db_path = 'D:\\Qeraat\\QeraatFasrhTools\\QeraatSearch\\qeraat_data_simple.db'
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Fetch all suras and their ayat count from the quran_sora table, ordered by sora number
cursor.execute('SELECT sora, sora_name, ayat_number FROM quran_sora ORDER BY sora')
soras = cursor.fetchall()

# Create a Word document
doc = Document()

# Add a title to the document
doc.add_heading(f'{site}', level=1)

# Function to download the content of a webpage
def download_content(sora, aya):
    url = f'https://tafsir.app/{site}/{sora}/{aya}'
    response = requests.get(url)
    if response.status_code == 200:
        # Use BeautifulSoup to parse the page content
        soup = BeautifulSoup(response.text, 'html.parser')
        
        # Find the div with id="preloaded-text"
        preloaded_text_div = soup.find('div', id='preloaded-text')
        if preloaded_text_div:
            # Extract the text content
            content = preloaded_text_div.get_text(strip=True)
            return content
    return None

# Iterate through each sora and its ayat
for sora, sora_name, ayat_number in soras:
    # Add the sora_name as a header for the sora
    doc.add_heading(f'سورة {sora_name}', level=2)
    
    # Loop through each aya in the sora
    for aya in range(1, ayat_number + 1):
        content = download_content(sora, aya)
        
        if content:
            # Add sora and aya as a heading
            doc.add_heading(f'آية {aya}', level=3)
            
            # Add the content to the document
            doc.add_paragraph(content)
            print(f'Successfully added content for سورة {sora_name}, آية {aya}')
        else:
            print(f'No content found for سورة {sora_name}, آية {aya}')

# Save the document
doc_path = 'D:\\Qeraat\\ibn_atiyah_content.docx'
doc.save(doc_path)
print(f'Document saved at {doc_path}')

# Close the database connection
conn.close()
