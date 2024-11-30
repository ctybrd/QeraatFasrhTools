import docx

# File paths
txt_file_path = r"e:\Qeraat\QeraatFasrhTools\quran_with_breaks.txt"
output_docx_path = r"e:\Qeraat\QeraatFasrhTools\quran_with_breaks.docx"

# Read the input text file
with open(txt_file_path, 'r', encoding='utf-8') as file:
    input_text = file.read()

# Split the text by form feed characters to separate pages
pages = input_text.split('\f')

# Create a new Word document
doc = docx.Document()

# Set the font style
style = doc.styles['Normal']
font = style.font
font.name = 'KFGQPC HAFS Uthmanic Script'
font.size = docx.shared.Pt(22)

# Add text for each page
for page_index, page in enumerate(pages):
    lines = page.split('\n')  # Split the page into lines

    for line in lines:
        # Add a paragraph with justified alignment
        para = doc.add_paragraph(line)
        para.alignment = docx.enum.text.WD_ALIGN_PARAGRAPH.JUSTIFY

    # If there are more pages, insert a page break after the current page (except the last one)
    if page_index < len(pages) - 1:
        doc.add_paragraph().add_run().add_break(docx.enum.text.WD_BREAK.PAGE)

# Save the document
doc.save(output_docx_path)

print(f"Word document saved to: {output_docx_path}")
