import sqlite3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt, Cm
from docx.shared import Cm

# Connect to the SQLite database
def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
        '(': ')',
        ')': '('
    }
    return ''.join(mapping.get(char, char) for char in str(number))


db_path = "E:/Qeraat/QeraatFasrhTools/QeraatSearch/Motshabeh/motshabeh.db"
connection = sqlite3.connect(db_path)
cursor = connection.cursor()

# Execute the query
query = """
    SELECT sora_name1, aya1, text1, sora_name2, aya2, text2
    FROM MotshabehatU
    WHERE sora_name1='البقرة' AND sora_name2='البقرة'
    ORDER BY aya1, aya2
"""
cursor.execute(query)
results = cursor.fetchall()

# Create a Word document
doc = Document()
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.7)
    section.left_margin = Cm(0.7)
    section.right_margin = Cm(0.7)

# Set font size for the entire document
doc.styles['Normal'].font.size = Pt(13)

# Function to add a paragraph with colored text
def add_colored_paragraph(doc, text, color):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.color.rgb = color

# Function to colorize common words in green
def colorize_common_words(paragraph, text1, text2):
    common_words = set(text1.split())
    current_run = paragraph.add_run()  # Create an initial run
    for word in text2.split():
        if current_run.text:
                current_run = paragraph.add_run(' ')
        if word in common_words:
            current_run = paragraph.add_run(word)
            current_run.font.color.rgb = RGBColor(255, 0, 0)  # Green
        else:
            # Append a space if there was a previous run
            current_run = paragraph.add_run(word)

# Iterate through the results and add paragraphs to the document
paragraph = doc.add_paragraph()
run = paragraph.add_run("متشابهات سورة البقرة\n")
run.font.color.rgb = RGBColor(0, 0, 0)

current_aya1 = None

for row in results:
    sora_name1, aya1, text1, sora_name2, aya2, text2 = row

    taya1 = transliterate_number(aya1)
    taya2 = transliterate_number(aya2)

    if current_aya1 != aya1:
        run = paragraph.add_run("=" * 50 + "\n")
        current_aya1 = aya1
        run = paragraph.add_run(f'{taya1} -  {text1}\n')
        run.font.color.rgb = RGBColor(0, 0, 255)  # Red

    # Colorize common words between text1 and text2 in green
    run = paragraph.add_run(f'{taya2} -')
    colorize_common_words(paragraph, text1, text2)
    run = paragraph.add_run('\n')
    # add_colored_paragraph(doc, f'{taya2} - {text2}\n', RGBColor(0, 0, 255))  # Blue

for paragraph in doc.paragraphs:
    paragraph.space_after = Pt(0)
    paragraph.space_before = Pt(0)

# Save the document
doc.save('Motshabehat70.docx')

# Close the database connection
connection.close()
