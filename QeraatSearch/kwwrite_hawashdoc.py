import sqlite3
import json
from docx import Document
from docx.shared import RGBColor

def cmyk_to_rgb(c, m, y, k):
    """Convert CMYK values to RGB."""
    r = 255 * (1.0 - c) * (1.0 - k)
    g = 255 * (1.0 - m) * (1.0 - k)
    b = 255 * (1.0 - y) * (1.0 - k)
    return int(r), int(g), int(b)

def get_data_from_db(db_path):
    """Retrieve data from the SQLite database."""
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()

    # Query to retrieve data grouped by Hawamesh_id
    cursor.execute('''
        SELECT Hawamesh_id, unicode, color, hawamesh.pagenumber, hawamesh.order1 
        FROM hawamesh 
        LEFT JOIN hawamesh_chars ON hawamesh.id = Hawamesh_chars.Hawamesh_id 
        ORDER BY hawamesh_chars.id
    ''')

    # Fetch all rows
    rows = cursor.fetchall()

    cursor.close()
    conn.close()

    return rows

def create_word_document(data, output_path):
    """Create a Word document with colored text based on the retrieved data."""
    document = Document()
    current_hawamesh_id = None
    paragraph = None

    for hawamesh_id, unicode_char, color_json, pageumber, order1 in data:
        # Check if we need to start a new paragraph
        if hawamesh_id != current_hawamesh_id:
            paragraph = document.add_paragraph()
            current_hawamesh_id = hawamesh_id
        
        # Parse the color JSON
        color_data = json.loads(color_json)
        c, m, y, k = color_data['ncolor']

        # Convert CMYK to RGB
        r, g, b = cmyk_to_rgb(c, m, y, k)

        # Add the unicode character to the paragraph
        run = paragraph.add_run(unicode_char)
        run.font.color.rgb = RGBColor(r, g, b)

    # Save the document
    document.save(output_path)

def main():
    db_path = 'd:/Qeraat/QeraatFasrhTools/QeraatSearch/kw.db'
    output_path = 'hawamesh_output.docx'

    # Get data from the database
    data = get_data_from_db(db_path)

    # Create a Word document with the data
    create_word_document(data, output_path)

if __name__ == '__main__':
    main()
