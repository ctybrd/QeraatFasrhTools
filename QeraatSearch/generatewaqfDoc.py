import sqlite3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Connect to your SQLite database
conn = sqlite3.connect('E:/Qeraat/QeraatFasrhTools/QeraatSearch/qeraat_data.db')
cursor = conn.cursor()

# Your SQL query
sql_query = """
SELECT
    qd.page_number2,
    wt.waqf AS waqf,
    GROUP_CONCAT('' || qd.aya || ' ـ ' || qd.sub_subject || '') AS aya_and_sub_subject
FROM
    quran_data qd
JOIN
    Waqf_Types wt ON 
    SUBSTR(qd.reading, INSTR(qd.reading, 'وقف ب')) = wt.waqf
WHERE qd.qarees LIKE '%حمزة%'
GROUP BY
    qd.page_number2, wt.waqf
ORDER BY
    qd.page_number2, wt.waqf
"""

# Execute the query
cursor.execute(sql_query)
result = cursor.fetchall()

# Create a Word document
doc = Document()

# Function to add aya_and_sub_subject with green and large font
def transliterate_number(text):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
    }
    return ''.join(mapping.get(char, char) for char in str(text))

def add_aya_and_sub_subject(paragraph, text):
    # Replace '(' with 'ـ' and remove ')'
    cleaned_text = text.replace('(', '').replace(')', '')
    run = paragraph.add_run(transliterate_number(cleaned_text))
    font = run.font
    font.size = Pt(14)
    font.color.rgb = RGBColor(0, 128, 0)  # Green color

# Variables to track the current page_number2
current_page_number2 = None

# Add header with page number
header = doc.sections[0].header
paragraph = header.paragraphs[0]
page_number_run = paragraph.add_run()
page_number_run.bold = True
paragraph.add_run("\n\n")  # Add line break

# Iterate through the query result
for row in result:
    # Check if the page_number2 has changed
    if row[0] != current_page_number2:
        # Add page_number2 and waqf
        doc.add_paragraph(f"صفحة : {transliterate_number(row[0])}", style='Heading1').alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Update the current_page_number2
        current_page_number2 = row[0]

    # Add aya_and_sub_subject with green color and large font
    aya_and_sub_subject_paragraph = doc.add_paragraph()
    add_aya_and_sub_subject(aya_and_sub_subject_paragraph, row[2])
    # Add waqf
    waqf = row[1]
    waqf = waqf.replace('.','')
    aya_and_sub_subject_paragraph.add_run(f" : {waqf}")

# Save the document
doc.save('output_document.docx')

# Close the database connection
conn.close()
