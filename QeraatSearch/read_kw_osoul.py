import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement, ns

def cmyk_to_rgb(c, m, y, k):
    """Convert CMYK values (0 to 1 range) to RGB values (0 to 255 range)."""
    r = 255 * (1 - c) * (1 - k)
    g = 255 * (1 - m) * (1 - k)
    b = 255 * (1 - y) * (1 - k)
    return int(r), int(g), int(b)

def process_detail_chars(detail_chars, cell):
    """Processes detail characters and formats them in a table cell."""
    paragraph = cell.paragraphs[0]
    for char_data in detail_chars:
        unicode_value = char_data.get('unicode', '')
        
        # Handle unicode values
        if unicode_value.startswith('0x'):
            try:
                text = chr(int(unicode_value, 16))
            except ValueError:
                text = ''
        else:
            text = unicode_value
        text = text.replace('s',' ')
        
        if text:
            run = paragraph.add_run(text)

            # Set font name
            fontname = char_data.get('fontname')
            if fontname:
                run.font.name = fontname
                rPr = run._element.get_or_add_rPr()
                rFonts = rPr.get_or_add_rFonts()
                rFonts.set(ns.qn('w:ascii'), fontname)
                rFonts.set(ns.qn('w:hAnsi'), fontname)
                rFonts.set(ns.qn('w:cs'), fontname)

            # Set font color
            color = char_data.get('color')
            if isinstance(color, dict) and 'ncolor' in color:
                cmyk_values = color['ncolor']
                if len(cmyk_values) == 4:
                    r, g, b = cmyk_to_rgb(*cmyk_values)
                    run.font.color.rgb = RGBColor(r, g, b)

            # Set font size (convert to Pt)
            size = char_data.get('size')
            if size:
                run.font.size = Pt(size)

            # Handle text direction if 'upright' is specified
            rtl = char_data.get('add_tab', False)
            if rtl:
                paragraph.add_run(' ' * 2)  # Add a tab space if specified
            rPr = run._element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.text = '1'  # '1' for right-to-left, '0' for left-to-right
            rPr.append(rtl)

def process_node(node, doc, distinct_key_chars):
    """Recursively processes each node and its child nodes, adding content to the document."""
    if isinstance(node, dict):
        key_chars = node.get('key_chars', [])
        value_chars = node.get('value_chars', [])

        # Record distinct key_chars entries as full strings
        key_chars_string = ""
        for char_data in key_chars:
            unicode_value = char_data.get('unicode', '')
            if unicode_value.startswith('0x'):
                try:
                    text = chr(int(unicode_value, 16))
                except ValueError:
                    text = ''
            else:
                text = unicode_value
            text = text.replace('s',' ')
            key_chars_string += text

        # Add the full key_chars string to the distinct set
        if key_chars_string:
            distinct_key_chars.add(key_chars_string)
        
        if key_chars or value_chars:
            # Create a table with two columns: one for key_chars and one for value_chars
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            key_cell = table.cell(0, 1)
            value_cell = table.cell(0, 0)
            
            if key_chars:
                process_detail_chars(key_chars, key_cell)
            
            if value_chars:
                process_detail_chars(value_chars, value_cell)

        # Recursively process nested nodes
        for key, value in node.items():
            if isinstance(value, list):
                for item in value:
                    process_node(item, doc, distinct_key_chars)

def insert_data(data, pagenumber, doc, distinct_key_chars):
    """Insert formatted data into the Word document."""
    doc.add_paragraph(f"صفحة: {pagenumber}").bold = True

    # Process each root node in the data
    for item in data:
        process_node(item, doc, distinct_key_chars)

def main():
    """Reads JSON data from multiple files in a folder and generates a Word document with the formatted text."""
    script_path = os.path.abspath(__file__)
    drive, _ = os.path.splitdrive(script_path)
    drive = drive + '/'
    folder_path = os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json')

    # Create a new Word document
    doc = Document()
    distinct_key_chars = set()  # To keep track of distinct key_chars

    try:
        # Extract numeric parts of filenames and sort them
        files = sorted(
            [f for f in os.listdir(folder_path) if f.startswith("Osoul_") and f.endswith(".json")],
            key=lambda x: int(x.split("_")[1].split(".")[0])
        )

        for filename in files:
            pagenumber = int(filename.split("_")[1].split(".")[0])  # Extract pagenumber from filename
            with open(os.path.join(folder_path, filename), 'r', encoding='utf-8') as f:
                data = json.load(f)
                insert_data(data, pagenumber, doc, distinct_key_chars)

                # Insert a page break after processing each file
                doc.add_page_break()

        # Save the Word document
        doc.save('Osoul_Content.docx')

        # Print the distinct key_chars entries
        print("Distinct key_chars entries:")
        for entry in distinct_key_chars:
            print(entry)

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == '__main__':
    main()
