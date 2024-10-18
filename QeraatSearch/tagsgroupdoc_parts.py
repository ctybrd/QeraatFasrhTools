import sqlite3
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Cm, RGBColor
from docx.oxml import OxmlElement

# Connect to the SQLite database
conn = sqlite3.connect('D:\\Qeraat\\QeraatFasrhTools\\QeraatSearch\\qeraat_data_simple.db')
cursor = conn.cursor()

# SQL query (remains unchanged)
query = """
WITH RECURSIVE tags_split(aya_index, id, sub_subject, reading, qareesrest, tag, remaining_tags) AS (
    SELECT aya_index, id, sub_subject, reading, qareesrest,
        CASE WHEN tags LIKE ',%' THEN SUBSTR(tags, 2, INSTR(tags, ',') - 1)
        ELSE SUBSTR(tags, 1, INSTR(tags, ',') - 1) END,
        CASE WHEN tags LIKE ',%' THEN SUBSTR(tags, INSTR(tags, ',') + 1)
        ELSE '' END
    FROM quran_data 
    WHERE tags IS NOT NULL AND tags NOT LIKE '%nochange%' AND page_shmrly IS NOT NULL

    UNION ALL

    SELECT aya_index, id, sub_subject, reading, qareesrest,
        CASE WHEN remaining_tags LIKE ',%' THEN SUBSTR(remaining_tags, 2, INSTR(remaining_tags, ',') - 1)
        ELSE SUBSTR(remaining_tags, 1, INSTR(remaining_tags, ',') - 1) END,
        CASE WHEN remaining_tags LIKE ',%' THEN SUBSTR(remaining_tags, INSTR(remaining_tags, ',') + 1)
        ELSE '' END
    FROM tags_split
    WHERE remaining_tags != ''
)

SELECT ts.aya_index, 
       ts.id, 
       ts.sub_subject, 
       ts.reading, 
       ts.tag, 
       tm.description,
       tm.category,
       qd.page_shmrly,
       qd.sora,
       qd.aya, qd.qareesrest
FROM tags_split ts
LEFT JOIN tagsmaster tm ON ts.tag = tm.tag
JOIN quran_data qd ON ts.aya_index = qd.aya_index AND ts.id = qd.id
WHERE ts.tag != '' and ts.tag !='meemsela' and ts.tag !='waqfhesham' and ts.tag !='waqfhamza'
ORDER BY qd.page_shmrly, ts.tag, ts.aya_index, ts.id;
"""

# Fetch data (remains unchanged)
data = cursor.execute(query).fetchall()
conn.close()

# Group data by page_shmrly (remains unchanged)
grouped_data = {}
for row in data:
    page_shmrly = row[7]  # page_shmrly column
    tag = row[4]           # Tag value
    if page_shmrly not in grouped_data:
        grouped_data[page_shmrly] = {}
    if tag not in grouped_data[page_shmrly]:
        grouped_data[page_shmrly][tag] = []
    grouped_data[page_shmrly][tag].append(row)

# Create Word documents in batches of 50 pages (remains unchanged)
doc = None
current_page_batch = 0

for page, tags in grouped_data.items():
    if current_page_batch % 50 == 0:  # New document for every 50 pages
        if doc:
            doc.save(f'shmrly_osoul_{current_page_batch // 50}.docx')  # Save previous document
        doc = Document()

        # Set page margins to 0.7 cm
        section = doc.sections[0]
        section.top_margin = Cm(0.7)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Cm(0.7)
        section.right_margin = Cm(0.7)

    # Add a page header with the page number
    doc.add_heading(f'صفحة: {page}', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for tag, records in tags.items():
        tag_description = records[0][5]  # Tag description
        category = records[0][6] or ''    # Set category to empty string if None
        doc.add_heading(f' {tag_description} - {category}', level=2)
        
        # Create a table for the records under the tag
        table = doc.add_table(rows=1, cols=4)  # 4 columns for aya, sub_subject, reading, qareesrest
        
        # Set column widths
        for i in range(1):  # Adjust widths for the first three columns
            table.columns[i].width = int(0.5 * 914400)  # 1.5 cm in EMU (1 cm = 914400 EMU)
        table.columns[1].width = int(3.0 * 914400)  # 3 cm for Sub Subject
        table.columns[2].width = int(5.0 * 914400)  # 5 cm for Reading
        table.columns[3].width = int(3.0 * 914400)  # 3 cm for Qareesrest
        
        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(record[9])  # Aya
            cell_sub_subject = row_cells[1]
            cell_sub_subject.text = str(record[2])  # Sub Subject
            
            # Change the text color of the sub_subject cell to green
            for paragraph in cell_sub_subject.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 128, 0)  # Green color
            
            row_cells[2].text = str(record[3])  # Reading
            row_cells[3].text = str(record[10])  # qareesrest
            
            # Reduce the space between rows
            for cell in row_cells:
                paragraph = cell.paragraphs[0]
                paragraph_format = paragraph.paragraph_format
                paragraph_format.space_after = Cm(0)  # No space after each paragraph
                paragraph_format.space_before = Cm(0)  # No space before each paragraph

                # Add RTL support directly to the paragraph element
                p_element = paragraph._element
                rPr = p_element.get_or_add_pPr()  # Get or add paragraph properties
                rtl = OxmlElement('w:rtl')
                rtl.text = '1'  # '1' for right-to-left
                rPr.append(rtl)

        doc.add_paragraph()  # Add a space after each tag section
    
    current_page_batch += 1
    doc.add_page_break()  # Add a page break after each page group

# Save the last document if it exists
if doc:
    doc.save(f'shmrly_osoul_{current_page_batch // 50}.docx')
