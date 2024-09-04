import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement, ns

def cmyk_to_rgb(c, m, y, k):
    """Convert CMYK values (0 to 1 range) to RGB values (0 to 255 range)."""
    r = 255 * (1 - c) * (1 - k)
    g = 255 * (1 - m) * (1 - k)
    b = 255 * (1 - y) * (1 - k)
    return int(r), int(g), int(b)

def process_qeraat(data, doc):
    """Process Qeraat JSON data."""
    for item in data:
        word_text = item.get('word_text', '')
        word_text = word_text.replace(')', '<b>').replace('(', '</b>')
        word_text = word_text.replace('<b>', '(').replace('</b>', ')')
        word_text = word_text.replace('﴿', '(').replace('﴾', ')')
        # word_text = word_text.replace('<b>', '﴾').replace('</b>', '﴿')
        word_text = word_text.replace('-', 'ـ')

        paragraph = doc.add_paragraph()  # Start a new paragraph
        paragraph.add_run(word_text).bold = True
        process_detail_chars(item.get('detail_title_chars', []), doc)
        for mp3_section in item.get('quran_ten_word_mp3', []):
            process_detail_chars(mp3_section.get('detail_chars', []), doc)

def process_osoul(data, doc):
    """Process Osoul JSON data."""
    distinct_key_chars = set()  # To keep track of distinct key_chars
    for item in data:
        process_node(item, doc, distinct_key_chars)

def process_shawahed(data, doc):
    """Process Shawahed JSON data."""
    for item in data:
        paragraph = doc.add_paragraph()  # Initialize paragraph at the start of each item
        for char_list in item.get('shahed_chars', []):
            for char_data in char_list:
                process_char_data(char_data, paragraph)
        for char_list in item.get('dalal_chars', []):
            for char_data in char_list:
                process_char_data(char_data, paragraph)

def process_hawamesh(data, doc):
    """Process Hawamesh JSON data."""
    for item in data:
        paragraph = doc.add_paragraph()  # Start a new paragraph
        for char_data in item.get('hawamesh_chars', []):
            text = char_data.get('unicode')
            if text == "*" or text[0] == '-':  # Break before '*'
                paragraph = doc.add_paragraph()  # Start a new paragraph
            text = text.replace(')', '<b>').replace('(', '</b>')
            text = text.replace('<b>', '(').replace('</b>', ')')
            text = text.replace('﴿', '(').replace('﴾', ')')
            # text = text.replace('<b>', '﴾').replace('</b>', '﴿')
            text = text.replace('-', 'ـ')
            run = paragraph.add_run(text)
            set_run_properties(run, char_data)

def set_run_properties(run, char_data):
    """Set properties such as font name, color, size, and direction for a run."""
    fontname = char_data.get('fontname')
    if fontname:
        run.font.name = fontname
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.get_or_add_rFonts()
        rFonts.set(ns.qn('w:ascii'), fontname)
        rFonts.set(ns.qn('w:hAnsi'), fontname)
        rFonts.set(ns.qn('w:cs'), fontname)

    color = char_data.get('color')
    if isinstance(color, dict) and 'ncolor' in color:
        cmyk_values = color['ncolor']
        if len(cmyk_values) == 4:
            r, g, b = cmyk_to_rgb(*cmyk_values)
            run.font.color.rgb = RGBColor(r, g, b)

    size = char_data.get('size')
    if size:
        run.font.size = Pt(size)

    if char_data.get('upright'):
        rPr = run._element.get_or_add_rPr()
        rtl = OxmlElement('w:rtl')
        rtl.text = '1'  # '1' for right-to-left, '0' for left-to-right
        rPr.append(rtl)

def process_detail_chars(detail_chars, doc):
    """Process and format detail characters."""
    paragraph = doc.add_paragraph()  # Start a new paragraph
    for char_data in detail_chars:
        text = char_data.get('unicode')
        text = text.replace('s', ' ')  # Example replacement
        text = text.replace(')', '<b>').replace('(', '</b>')
        text = text.replace('<b>', '(').replace('</b>', ')')
        text = text.replace('-', 'ـ')

        run = paragraph.add_run(text)
        set_run_properties(run, char_data)

def process_node(node, doc, distinct_key_chars):
    """Recursively process each node for Osoul and add content to the document."""
    if isinstance(node, dict):
        key_chars = node.get('key_chars', [])
        value_chars = node.get('value_chars', [])

        key_chars_string = "".join([chr(int(char_data['unicode'], 16)) if char_data['unicode'].startswith('0x') else char_data['unicode'] for char_data in key_chars])
        if key_chars_string:
            distinct_key_chars.add(key_chars_string)
        
        if key_chars or value_chars:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            key_cell = table.cell(0, 1)
            value_cell = table.cell(0, 0)
            process_detail_chars(key_chars, key_cell)
            process_detail_chars(value_chars, value_cell)

        for key, value in node.items():
            if isinstance(value, list):
                for item in value:
                    process_node(item, doc, distinct_key_chars)

def process_char_data(char_data, paragraph):
    """Process individual character data for Shawahed and add it to the document."""
    text = char_data.get('unicode')
    if not text:
        return
    if text == "*" or text[0] == '-':  # Break before '*'
        paragraph = doc.add_paragraph()  # Start a new paragraph
    text = text.replace(')', '<b>').replace('(', '</b>')
    text = text.replace('<b>', '(').replace('</b>', ')')
    text = text.replace('-', 'ـ')
    run = paragraph.add_run(text)
    set_run_properties(run, char_data)
    if text.endswith('.'):
        paragraph.add_run().add_break()

def main():
    """Main function to process all JSON formats and combine them into one Word document."""
    script_path = os.path.abspath(__file__)
    drive, _ = os.path.splitdrive(script_path)
    drive = drive + '/'
    folders = {
        "qeraat": os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json'),
        "osoul": os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json'),
        "shawahed": os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json'),
        "hawamesh": os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json')
    }

    # Create a new Word document
    doc = Document()

    try:
        max_pages = 0
        file_groups = {}

        # Load and sort files for each type
        for key, folder_path in folders.items():
            files = sorted(
                [f for f in os.listdir(folder_path) if f.startswith(key.capitalize()) and f.endswith(".json")],
                key=lambda x: int(x.split("_")[1].split(".")[0])
            )
            max_pages = max(max_pages, len(files))
            file_groups[key] = files

        # Process files page by page
        for i in range(max_pages):
            if i < len(file_groups["qeraat"]):
                filename = file_groups["qeraat"][i]
                pagenumber = int(filename.split("_")[1].split(".")[0])
                with open(os.path.join(folders["qeraat"], filename), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    doc.add_paragraph(f"Qeraat - صفحة: {pagenumber}").bold = True
                    process_qeraat(data, doc)

            if i < len(file_groups["osoul"]):
                filename = file_groups["osoul"][i]
                pagenumber = int(filename.split("_")[1].split(".")[0])
                with open(os.path.join(folders["osoul"], filename), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    doc.add_paragraph(f"Osoul - صفحة: {pagenumber}").bold = True
                    process_osoul(data, doc)

            if i < len(file_groups["shawahed"]):
                filename = file_groups["shawahed"][i]
                pagenumber = int(filename.split("_")[1].split(".")[0])
                with open(os.path.join(folders["shawahed"], filename), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    doc.add_paragraph(f"Shawahed - صفحة: {pagenumber}").bold = True
                    process_shawahed(data, doc)

            if i < len(file_groups["hawamesh"]):
                filename = file_groups["hawamesh"][i]
                pagenumber = int(filename.split("_")[1].split(".")[0])
                with open(os.path.join(folders["hawamesh"], filename), 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    doc.add_paragraph(f"Hawamesh - صفحة: {pagenumber}").bold = True
                    process_hawamesh(data, doc)

            doc.add_page_break()

        # Save the combined document
        doc.save(os.path.join(drive, "c.docx"))

    except Exception as e:
        print("Error processing files:", e)

if __name__ == "__main__":
    main()
