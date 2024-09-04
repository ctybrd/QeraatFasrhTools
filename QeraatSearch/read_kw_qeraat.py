import json
import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement, ns

def cmyk_to_rgb(c, m, y, k):
    """Convert CMYK values (0 to 1 range) to RGB values (0 to 255 range)."""
    r = 255 * (1 - c) * (1 - k)
    g = 255 * (1 - m) * (1 - k)
    b = 255 * (1 - y) * (1 - k)
    return int(r), int(g), int(b)

def insert_data(data, pagenumber, doc):
    """Formats the JSON data in a Word document.

    Args:
        data: The JSON data to format.
        pagenumber: The page number associated with the data.
        doc: The Word document object to add the content to.
    """
    doc.add_paragraph(f"صفحة: {pagenumber}").bold = True

    for item in data:
        word_text = item.get('word_text', '')
        paragraph = doc.add_paragraph()  # Start a new paragraph
        paragraph.add_run(word_text).bold = True

        # Process `detail_title_chars`
        process_detail_chars(item.get('detail_title_chars', []), doc)

        # Process `quran_ten_word_mp3`
        for mp3_section in item.get('quran_ten_word_mp3', []):
            process_detail_chars(mp3_section.get('detail_chars', []), doc)

def process_detail_chars(detail_chars, doc):
    """Processes detail characters and adds them to the Word document."""
    paragraph = doc.add_paragraph()  # Start a new paragraph
    for char_data in detail_chars:
        text = char_data.get('unicode')
        text = text.replace('s', ' ')  # Example replacement

        run = paragraph.add_run(text)

        # Set font name
        fontname = char_data.get('fontname')
        if fontname:
            run.font.name = fontname
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.get_or_add_rFonts()
            rFonts.set(ns.qn('w:ascii'), fontname)
            rFonts.set(ns.qn('w:hAnsi'), fontname)
            rFonts.set(ns.qn('w:cs'), fontname)

        # Set font color
        color = char_data.get('color')
        if isinstance(color, dict) and 'ncolor' in color:
            cmyk_values = color['ncolor']
            if len(cmyk_values) == 4:
                r, g, b = cmyk_to_rgb(*cmyk_values)
                run.font.color.rgb = RGBColor(r, g, b)

        # Set font size (convert to Pt)
        size = char_data.get('size')
        if size:
            run.font.size = Pt(size)

        # Handle text direction if 'upright' is specified
        if char_data.get('upright'):
            rPr = run._element.get_or_add_rPr()
            rtl = OxmlElement('w:rtl')
            rtl.text = '1'  # '1' for right-to-left, '0' for left-to-right
            rPr.append(rtl)

def main():
    """Reads JSON data from multiple files in a folder and generates a Word document with the formatted text."""
    script_path = os.path.abspath(__file__)
    drive, _ = os.path.splitdrive(script_path)
    drive = drive + '/'
    folder_path = os.path.join(drive, 'Qeraat/QeraatFasrhTools_Data/Ten_Readings/json')

    # Create a new Word document
    doc = Document()

    try:
        # Extract numeric parts of filenames and sort them
        files = sorted(
            [f for f in os.listdir(folder_path) if f.startswith("Qeraat_") and f.endswith(".json")],
            key=lambda x: int(x.split("_")[1].split(".")[0])
        )

        for filename in files:
            pagenumber = int(filename.split("_")[1].split(".")[0])  # Extract pagenumber from filename
            with open(os.path.join(folder_path, filename), 'r', encoding='utf-8') as f:
                data = json.load(f)
                insert_data(data, pagenumber, doc)

                # Insert a page break after processing each file
                doc.add_page_break()

        # Save the Word document
        doc.save('qeraat_Content.docx')

    except Exception as e:
        print(f"An error occurred: {e}")

if __name__ == '__main__':
    main()
