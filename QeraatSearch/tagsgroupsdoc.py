import sqlite3
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Connect to the SQLite database
conn = sqlite3.connect('D:\\Qeraat\\QeraatFasrhTools\\QeraatSearch\\qeraat_data_simple.db')
cursor = conn.cursor()

# Execute the updated SQL query
query = """
WITH RECURSIVE tags_split(aya_index, id, sub_subject, reading, tag, remaining_tags) AS (
    SELECT aya_index, id, sub_subject, reading,
        CASE WHEN tags LIKE ',%' THEN SUBSTR(tags, 2, INSTR(tags, ',') - 1)
        ELSE SUBSTR(tags, 1, INSTR(tags, ',') - 1) END,
        CASE WHEN tags LIKE ',%' THEN SUBSTR(tags, INSTR(tags, ',') + 1)
        ELSE '' END
    FROM quran_data 
    WHERE tags IS NOT NULL AND tags NOT LIKE '%nochange%' AND page_shmrly IS NOT NULL

    UNION ALL

    SELECT aya_index, id, sub_subject, reading,
        CASE WHEN remaining_tags LIKE ',%' THEN SUBSTR(remaining_tags, 2, INSTR(remaining_tags, ',') - 1)
        ELSE SUBSTR(remaining_tags, 1, INSTR(remaining_tags, ',') - 1) END,
        CASE WHEN remaining_tags LIKE ',%' THEN SUBSTR(remaining_tags, INSTR(remaining_tags, ',') + 1)
        ELSE '' END
    FROM tags_split
    WHERE remaining_tags != ''
)

SELECT ts.aya_index, 
       ts.id, 
       ts.sub_subject, 
       ts.reading, 
       ts.tag, 
       tm.description,
       tm.category,
       qd.page_shmrly,
       qd.sora,
       qd.aya
FROM tags_split ts
LEFT JOIN tagsmaster tm ON ts.tag = tm.tag
JOIN quran_data qd ON ts.aya_index = qd.aya_index AND ts.id = qd.id
WHERE ts.tag != ''
ORDER BY qd.page_shmrly, ts.tag, ts.aya_index, ts.id;
"""

# Fetch data
data = cursor.execute(query).fetchall()
conn.close()

# Create a new Word document
doc = Document()

# Group data by page_shmrly and then by tag
grouped_data = {}
for row in data:
    page_shmrly = row[-2]  # Page identifier
    tag = row[4]           # Tag value
    if page_shmrly not in grouped_data:
        grouped_data[page_shmrly] = {}
    if tag not in grouped_data[page_shmrly]:
        grouped_data[page_shmrly][tag] = []
    grouped_data[page_shmrly][tag].append(row)

# Generate the document
for page, tags in grouped_data.items():
    # Add a page header
    doc.add_heading(f'صفحة: {page}', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    for tag, records in tags.items():
        tag_description = records[0][5]  # Tag description
        category = records[0][6] or ''    # Set category to '' if empty
        doc.add_heading(f' {tag_description} - {category}', level=2)
        
        # Create a table for the records under the tag
        table = doc.add_table(rows=1, cols=5)  # 5 columns for sora, aya, id, sub_subject, reading
        
        # Setting header values
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = ''  # No English title for Sora
        hdr_cells[1].text = ''  # No English title for Aya
        hdr_cells[2].text = ''  # No English title for ID
        hdr_cells[3].text = ''  # No English title for Sub Subject
        hdr_cells[4].text = ''  # No English title for Reading
        
        # Set column widths
        for i in range(3):  # Adjust widths for first three columns
            table.columns[i].width = int(1.5 * 914400)  # 1.5 cm in EMU (1 cm = 914400 EMU)
        table.columns[3].width = int(2.0 * 914400)  # 2 cm for Sub Subject
        table.columns[4].width = int(2.0 * 914400)  # 2 cm for Reading
        
        for record in records:
            row_cells = table.add_row().cells
            row_cells[0].text = str(record[8])  # Sora
            row_cells[1].text = str(record[9])  # Aya
            row_cells[2].text = str(record[1])  # ID
            row_cells[3].text = str(record[2])  # Sub Subject
            row_cells[4].text = str(record[3])  # Reading

        doc.add_paragraph()  # Add a space after each tag section
    
    doc.add_page_break()  # Add a page break after each page

# Save the document
doc.save('output.docx')
