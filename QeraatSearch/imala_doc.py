from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import sqlite3

# Connect to the SQLite database
db_path = "e:\\Qeraat\\QeraatFasrhTools\\QeraatSearch\\qeraat_data_simple.db"
conn = sqlite3.connect(db_path)
cursor = conn.cursor()

# Define the SQL query
# sql_query = """

# WITH summed_data AS (
#     SELECT 
#         page_shmrly,
#         CASE WHEN tags LIKE '%imala%' THEN 
# 		'الإمالة' ELSE 
# 		'التقليل' 
# 		END AS description,
#         CASE WHEN reading LIKE '%بخلف%' THEN ' ـ بخلف عنه ـ' ELSE '' END AS kholf,
#         IFNULL(resultnew, sub_subject1) AS sub_subject,
#         --CASE WHEN rasaya IS NOT NULL THEN 'رؤوس الآي' ELSE 'ما ليس برأس آية' END AS rasaya,
# 		'' as rasaya,
#         SUM(q1) AS q1, SUM(r1_1) AS r1_1, SUM(r1_2) AS r1_2,
#         SUM(q2) AS q2, SUM(r2_1) AS r2_1, SUM(r2_2) AS r2_2,
#         SUM(q3) AS q3, SUM(r3_1) AS r3_1, SUM(r3_2) AS r3_2,
#         SUM(q4) AS q4, SUM(r4_1) AS r4_1, SUM(r4_2) AS r4_2,
#         SUM(q5) AS q5, SUM(r5_1) AS r5_1, SUM(r5_2) AS r5_2,
#         SUM(q6) AS q6, SUM(r6_1) AS r6_1, SUM(r6_2) AS r6_2,
#         SUM(q7) AS q7, SUM(r7_1) AS r7_1, SUM(r7_2) AS r7_2,
#         SUM(q8) AS q8, SUM(r8_1) AS r8_1, SUM(r8_2) AS r8_2,
#         SUM(q9) AS q9, SUM(r9_1) AS r9_1, SUM(r9_2) AS r9_2,
#         SUM(q10) AS q10, SUM(r10_1) AS r10_1, SUM(r10_2) AS r10_2
#     FROM quran_data
#     WHERE tags LIKE '%,imala,%' OR tags LIKE '%,taklel,%'
#     GROUP BY page_shmrly, description, kholf, sub_subject, rasaya
# )
# SELECT page_shmrly,description,kholf,group_concat(' ﴾'|| sub_subject || '﴿ ') sub_subject,rasaya,
#     TRIM(
#         CASE WHEN q1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q1') || ', ' ELSE '' END ||
#         CASE WHEN q1 IS NULL AND r1_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R1_1') || ', ' ELSE '' END ||
#         CASE WHEN q1 IS NULL AND r1_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R1_2') || ', ' ELSE '' END ||
#         CASE WHEN q2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q2') || ', ' ELSE '' END ||
#         CASE WHEN q2 IS NULL AND r2_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R2_1') || ', ' ELSE '' END ||
#         CASE WHEN q2 IS NULL AND r2_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R2_2') || ', ' ELSE '' END ||
#         CASE WHEN q3 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q3') || ', ' ELSE '' END ||
#         CASE WHEN q3 IS NULL AND r3_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R3_1') || ', ' ELSE '' END ||
#         CASE WHEN q3 IS NULL AND r3_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R3_2') || ', ' ELSE '' END ||
#         CASE WHEN q4 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q4') || ', ' ELSE '' END ||
#         CASE WHEN q4 IS NULL AND r4_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R4_1') || ', ' ELSE '' END ||
#         CASE WHEN q4 IS NULL AND r4_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R4_2') || ', ' ELSE '' END ||
#         CASE WHEN q5 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q5') || ', ' ELSE '' END ||
#         CASE WHEN q5 IS NULL AND r5_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R5_1') || ', ' ELSE '' END ||
#         CASE WHEN q5 IS NULL AND r5_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R5_2') || ', ' ELSE '' END ||
#         CASE WHEN q6 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q6') || ', ' ELSE '' END ||
#         CASE WHEN q6 IS NULL AND r6_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R6_1') || ', ' ELSE '' END ||
#         CASE WHEN q6 IS NULL AND r6_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R6_2') || ', ' ELSE '' END ||
#         CASE WHEN q7 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q7') || ', ' ELSE '' END ||
#         CASE WHEN q7 IS NULL AND r7_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R7_1') || ', ' ELSE '' END ||
#         CASE WHEN q7 IS NULL AND r7_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R7_2') || ', ' ELSE '' END ||
#         CASE WHEN q8 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q8') || ', ' ELSE '' END ||
#         CASE WHEN q8 IS NULL AND r8_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R8_1') || ', ' ELSE '' END ||
#         CASE WHEN q8 IS NULL AND r8_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R8_2') || ', ' ELSE '' END ||
#         CASE WHEN q9 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q9') || ', ' ELSE '' END ||
#         CASE WHEN q9 IS NULL AND r9_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R9_1') || ', ' ELSE '' END ||
#         CASE WHEN q9 IS NULL AND r9_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R9_2') || ', ' ELSE '' END ||
#         CASE WHEN q10 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'Q10') || ', ' ELSE '' END ||
#         CASE WHEN q10 IS NULL AND r10_1 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R10_1') || ', ' ELSE '' END ||
#         CASE WHEN q10 IS NULL AND r10_2 IS NOT NULL THEN (SELECT name FROM qareemaster WHERE qkey = 'R10_2') || ', ' ELSE '' END
#     , ', ') AS qarees
# FROM summed_data
# group by page_shmrly, description, kholf, rasaya,qarees
# order by page_shmrly,  rasaya,description, kholf;

# """
sql_query ="""

WITH summed_data AS (
    SELECT 
        page_shmrly,
		' إمالة هاء التأنيث وقفا (الكسائي)'
		AS description,
        CASE WHEN reading LIKE '%بخلف%' THEN ' ـ بخلف ـ' ELSE '' END AS kholf,
        IFNULL(resultnew, sub_subject1) AS sub_subject,

		'' as rasaya
    FROM quran_data
    WHERE tags LIKE '%,imalah,%' 
    GROUP BY page_shmrly, description, kholf, sub_subject, rasaya
	ORDER by aya_index,id
)
SELECT page_shmrly,description,kholf,group_concat(' ﴾'|| sub_subject || '﴿ ') sub_subject,rasaya,'' AS qarees
FROM summed_data
group by page_shmrly,kholf
order by page_shmrly,kholf

"""
# Execute the SQL query
cursor.execute(sql_query)
rows = cursor.fetchall()

# Close the database connection
conn.close()

# Initialize variables for document creation
doc = Document()
current_page_shmrly = None
current_description = None
file_count = 1

def save_and_create_new_doc():
    global doc, file_count
    doc.save(f'e:/Qeraat/مشروع العشر/shmrly_imalah_{file_count}.docx')
    file_count += 1
    doc = Document()

# Helper function to add formatted text
def add_text(paragraph, text, font_color, underline=False, bold=False, size=12):
    run = paragraph.add_run(text)
    run.font.color.rgb = RGBColor(*font_color)
    run.font.underline = underline
    run.font.bold = bold
    run.font.size = Pt(size)
    run.add_text(' ')
# Initialize variables for tracking the current group
current_page_shmrly = None
current_description = None
first_row_in_group = True

# Process each row
for row in rows:
    page_shmrly, description, kholf, sub_subject, rasaya, qarees = row

    # Add a header for a new page_shmrly
    if page_shmrly != current_page_shmrly:
        doc.add_heading(f'صفحة: {page_shmrly}', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        current_page_shmrly = page_shmrly
        current_description = None
        first_row_in_group = True

    # Start a new paragraph only if the description changes
    if description != current_description:
        para = doc.add_paragraph()
        current_description = description
        first_row_in_group = True

    # Add `/` at the beginning of the row if it's not the first row in the current group
    if not first_row_in_group:
        add_text(para, '/ ', font_color=(0, 0, 0))

    # Add text for each part with specific colors and suppress repeating description
    if first_row_in_group:
        add_text(para, description + ':', font_color=(210, 35, 41))  # Red color for description
        first_row_in_group = False

    add_text(para, kholf, font_color=(0, 0, 255))          # Blue color for kholf
    add_text(para, sub_subject, font_color=(0, 128, 0))    # Green color for sub_subject
    # add_text(para, qarees, font_color=(0, 0, 0))           # Black color for qarees
    

# Save the final document
save_and_create_new_doc()

"Document generation completed successfully."
