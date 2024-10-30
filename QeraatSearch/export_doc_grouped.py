import sqlite3
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩'
    }
    return ''.join(mapping.get(char, char) for char in str(number))

# Define color for serial numbers
RED = RGBColor(0, 0, 0)

# Open the SQLite database
db_path = "D:/Qeraat/QeraatFasrhTools/QeraatSearch/qeraat_data_simple.db"
connection = sqlite3.connect(db_path)
cursor = connection.cursor()

# Execute the updated SQL query
query = """
SELECT CASE SUBSTRING(sub_subject1, INSTR(sub_subject1, ' ') + 2, 1)
    WHEN 'َ' THEN 'المفتوح'
    WHEN 'ُ' THEN 'المضموم' 
    WHEN 'ِ' THEN 'المكسور' 
    END AS harakaname,
    IFNULL(resultnew, sub_subject1) AS sub_subject1,
    sora_name,
    aya
FROM all_qeraat
WHERE tags LIKE '%,ikhfaa,%'
    AND (r8_1 IS NOT NULL OR r8_2 IS NOT NULL)
ORDER BY SUBSTRING(sub_subject1, INSTR(sub_subject1, ' ') + 2, 1), aya_index, id
"""
cursor.execute(query)

# Create a new Word document
doc = Document()

# Add document title
title = doc.add_paragraph("حصر مواضع إخفاء أبي جعفر عند الغين والخاء")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.bold = True
title_run.font.size = Pt(18)

# Initialize serial number and current harakaname
serial_num = 1
current_harakaname = None

# Process each row from the query
for row in cursor.fetchall():
    harakaname, sub_subject1, sora_name, aya = row

    # Add new harakaname group header
    if harakaname != current_harakaname:
        current_harakaname = harakaname
        header = doc.add_paragraph(harakaname)
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        header_run = header.runs[0]
        header_run.bold = True
        header_run.font.size = Pt(16)  # Large font for the group header

    # Transliterated serial number
    serial_num_arabic = transliterate_number(serial_num)

    # Create a new paragraph for each entry
    entry = doc.add_paragraph()
    entry.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add serial number and sub_subject1
    serial_num_run = entry.add_run(f"{serial_num_arabic}ـ  ) {sub_subject1} (  ")
    serial_num_run.font.size = Pt(14)  # Font size for serial number and sub_subject1
    serial_num_run.font.color.rgb = RED

    # Add sora_name and aya with a smaller font
    sora_aya_run = entry.add_run(f"{sora_name}: {transliterate_number(aya)}")
    sora_aya_run.font.size = Pt(10)  # Smaller font size for sora_name and aya

    # Increment serial number for the next entry
    serial_num += 1

# Save the document
doc.save("./output/ikhfaa_grouped.docx")

# Close the database connection
connection.close()
