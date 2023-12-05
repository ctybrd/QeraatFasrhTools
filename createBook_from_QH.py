import sqlite3
from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.shared import RGBColor
from PIL import Image, ImageOps
import re
from docx.shared import Cm


def transliterate_number(number):
    mapping = {
        '0': '٠',
        '1': '١',
        '2': '٢',
        '3': '٣',
        '4': '٤',
        '5': '٥',
        '6': '٦',
        '7': '٧',
        '8': '٨',
        '9': '٩',
        '(': ')',
        ')': '(',
        '[': ']',
        ']': '[',
         
            
    }
    return ''.join(mapping.get(char, char) for char in str(number))

def connect_to_database(database_path):
    conn = sqlite3.connect(database_path)
    cursor = conn.cursor()
    return conn, cursor

def retrieve_data(cursor):
    query = '''
        select page_number,text,case when min(aya_number)=max(aya_number) then min(aya_number) else min(aya_number)||' - ' || max(aya_number) end as ayas1,
        group_concat(ayakey) gkey 
        from(Select printf('%03d', mosshf_shmrly.sora_number) || printf('%03d', mosshf_shmrly.aya_number)  as ayakey,book_moyassar.aya_index,book_moyassar.text,mosshf_shmrly.page_number,
        mosshf_shmrly.aya_number,mosshf_shmrly.sora_number from book_moyassar
        join mosshf_shmrly on book_moyassar.aya_index=mosshf_shmrly.aya_index
        order by book_moyassar.aya_index)
        group by page_number,text
        order by page_number,gkey
    '''
    cursor.execute(query)
    data = cursor.fetchall()
    return data

def process_html_text(html_text):
    soup = BeautifulSoup(html_text, 'html.parser')
    return soup.get_text()

def add_frame_to_image(image_path):
    image = Image.open(image_path)
    framed_image = ImageOps.expand(image, border=4, fill='grey')
    framed_image_path = 'framed_image.png'
    framed_image.save(framed_image_path)
    return framed_image_path

def create_word_document(data):
    doc = Document()

    # Set consistent page margins for all sections
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(0.7)
        section.bottom_margin = Cm(0.7)
        section.left_margin = Cm(0.7)
        section.right_margin = Cm(0.7)

# Set font size for the entire document
    doc.styles['Normal'].font.size = Pt(13)

    current_page = None
    color_switch = False  # Flag to switch between black and blue colors

    for page_number,text, aya_number,gkey in data:
        if current_page is None or current_page != page_number:
            if current_page is not None:
                doc.add_page_break()

            current_page = page_number

            image_path = f'e:/othmanc/{current_page+2}.png'  # Replace with the actual path to the folder and image file extension

            # Add frame to the image
            framed_image_path = add_frame_to_image(image_path)

            paragraph = doc.add_paragraph()
            if current_page % 2 == 0:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            else:
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = paragraph.add_run()
            run.add_picture(framed_image_path, width=Inches(4.5))


            comments_paragraph = doc.add_paragraph()
            comments_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT    

        processed_text = process_html_text(text)
        processed_text = processed_text.replace('.', '')
        processed_text = processed_text.replace('\r\n', ' ـ ')
        processed_text = processed_text.replace('\r', ' ـ ')
        processed_text = processed_text.replace('\n', ' ـ ')
        processed_text = transliterate_number( str(aya_number) + 'ـ ' + processed_text+'\n')
        # processed_text = re.sub(' {2,}', ' ', processed_text)

        
        run = comments_paragraph.add_run()

        # Set text color to black or blue based on the color_switch flag
        if color_switch:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # Blue color
        else:
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # Black color


        run.text = processed_text
        comments_paragraph.paragraph_format.line_spacing = Pt(13)

        color_switch = not color_switch  # Toggle the color_switch flag

    doc.save('output.docx')

def close_database_connection(connection):
    connection.close()

# Main code
database_path = 'e:/qeraat/data_v17.db'
image_folder = 'e:/othmanc/'

conn, cursor = connect_to_database(database_path)
data = retrieve_data(cursor)
create_word_document(data)
close_database_connection(conn)
